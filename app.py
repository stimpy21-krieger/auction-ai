import streamlit as st
import requests
import uuid
import time
import json
import hashlib
import re
import os
import pandas as pd
import datetime
import numpy as np
import cv2
import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image
import pillow_heif  # HEIC/HEIF 지원 (아이폰 사진)
import plotly.graph_objects as go
import typing
import asyncio
import aiohttp
from functools import lru_cache

# =====================================================================
# 🎨 1. 스트림릿 기본 설정 & 2030 타겟 감성 디자인 (CSS)
# =====================================================================
st.set_page_config(page_title="AI 경매 권리분석 마법사", page_icon="🧙‍♂️", layout="centered")

st.markdown("""
<style>
    /* 트렌디한 Pretendard 폰트 적용 및 부드러운 배경색 (이모지 폰트 최우선 적용) */
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    
    html, body, [class*="css"] {
        font-family: 'Apple Color Emoji', 'Segoe UI Emoji', 'Noto Color Emoji', 'Pretendard', -apple-system, sans-serif !important;
        background-color: #FDFBF7;
        color: #4A4A4A;
    }
    
    h1 { font-size: 26px !important; color: #333333; font-weight: 700; margin-bottom: 5px !important;}
    h2, h3 { font-size: 20px !important; color: #4A4A4A; font-weight: 600; }
    p, li, span { font-size: 15px !important; line-height: 1.6; }
    
    /* 코랄 핑크빛 둥근 메인 버튼 */
    .stButton > button[kind="primary"] {
        background-color: #FF9B9B !important;
        color: white !important;
        border-radius: 25px !important;
        border: none !important;
        padding: 10px 20px !important;
        font-size: 16px !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 10px rgba(255, 155, 155, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #FF8282 !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 14px rgba(255, 155, 155, 0.4) !important;
    }

    /* 🌟 파일 업로드 창 스타일링 🌟 */
    [data-testid="stFileUploadDropzone"] {
        background-color: #FFFFFF !important;
        border: 2px dashed #FF9B9B !important;
        border-radius: 20px !important;
        padding: 30px !important;
    }

    /* Browse files 버튼 — JS MutationObserver가 텍스트를 치환하므로 CSS ::after 불필요 */
    [data-testid="stFileUploadDropzone"] button {
        font-weight: 600 !important;
        font-size: 14px !important;
    }
</style>

<script>
// 🌟 파일 업로드 영어 텍스트 → 한글 강제 치환 (MutationObserver)
(function() {
    const REPLACEMENTS = {
        'Drag and drop file here': '📸 터치해서 등기부등본 사진 올리기',
        'Drag and drop files here': '📸 터치해서 등기부등본 사진 올리기',
        'Browse files': '사진 선택하기',
        'Browse file': '사진 선택하기'
    };
    // 이미 교체된 한글 문구 목록 (중복 교체 방지)
    const KO_VALUES = new Set(Object.values(REPLACEMENTS));
    KO_VALUES.add('파일당 최대 500MB');

    // 업로드 중 DOM 변경 방지 플래그
    let isUploading = false;
    // MutationObserver 재진입 방지 플래그
    let isReplacing = false;

    function replaceTexts() {
        // 업로드 중이거나, 이미 교체 작업 중이면 스킵
        if (isUploading || isReplacing) return;
        isReplacing = true;

        try {
            const dropzones = document.querySelectorAll('[data-testid="stFileUploadDropzone"]');
            dropzones.forEach(zone => {
                // 업로드 진행 중인 dropzone은 건드리지 않음 (progress bar 감지)
                if (zone.querySelector('[role="progressbar"]') ||
                    zone.querySelector('[data-testid="stFileUploadDeleteBtn"]')) {
                    return;
                }

                const walker = document.createTreeWalker(zone, NodeFilter.SHOW_TEXT, null, false);
                let node;
                while (node = walker.nextNode()) {
                    const text = node.textContent;
                    const trimmed = text.trim();

                    // 이미 한글로 교체된 텍스트는 스킵
                    if (KO_VALUES.has(trimmed)) continue;

                    // 정확히 매칭되는 영어 문구 교체
                    for (const [eng, kor] of Object.entries(REPLACEMENTS)) {
                        if (trimmed === eng) {
                            node.textContent = kor;
                            break;
                        }
                    }

                    // 용량 제한 문구 교체 (부분 매칭: "Limit 200MB per file · JPG, ..." 전체 처리)
                    const limitMatch = trimmed.match(/Limit\s+\d+(?:\.\d+)?\s*(?:MB|KB|GB)\s+per\s+file/i);
                    if (limitMatch) {
                        node.textContent = text.replace(limitMatch[0], '파일당 최대 500MB');
                    }
                }
            });
        } finally {
            isReplacing = false;
        }
    }

    // 파일 input 변경 감지 → 업로드 중 DOM 조작 완전 차단
    document.addEventListener('change', function(e) {
        if (e.target && e.target.type === 'file') {
            isUploading = true;
            // 업로드 완료 후 DOM 조작 재개 (대용량 다중 파일 고려하여 충분한 대기)
            setTimeout(function() {
                isUploading = false;
                replaceTexts();
            }, 5000);
        }
    }, true);

    // Debounce: 빠른 연속 호출 방지
    let debounceTimer = null;
    function debouncedReplaceTexts() {
        if (debounceTimer) clearTimeout(debounceTimer);
        debounceTimer = setTimeout(replaceTexts, 200);
    }

    // 초기 실행 + DOM 변경 감시 (characterData 포함 — Streamlit 텍스트 업데이트 감지 필수)
    const observer = new MutationObserver(debouncedReplaceTexts);
    observer.observe(document.body, { childList: true, subtree: true, characterData: true });
    // 페이지 로드 후 딜레이를 두고 실행 (Streamlit 렌더링 대기)
    setTimeout(replaceTexts, 300);
    setTimeout(replaceTexts, 1000);
    setTimeout(replaceTexts, 2000);
    setTimeout(replaceTexts, 4000);
})();
</script>
""", unsafe_allow_html=True)

# =====================================================================
# 🖼️ 스마트 이미지 전처리 (Grayscale + CLAHE + Deskew)
# =====================================================================
def smart_preprocess(file_bytes):
    """스마트폰 사진을 OCR 최적화 전처리합니다.
    - HEIC/HEIF → PNG 자동 변환 (아이폰 사진 지원)
    - PIL RGB 강제 변환 → 메타데이터 충돌 100% 방지
    - 해상도 유지 (resize 없음)
    - Grayscale + CLAHE 대비 향상 → 용량 축소 + 선명도 향상
    - Hough Line Transform 기반 deskew → 기울어진 문서 자동 수평 보정
    - PNG 무손실 출력
    """
    try:
        # ===== [1] HEIC/HEIF 변환 (아이폰 사진 지원) =====
        try:
            heif_image = pillow_heif.read_heif(file_bytes)
            pil_img = Image.frombytes(
                heif_image.mode, heif_image.size, heif_image.data,
                "raw", heif_image.mode, heif_image.stride
            )
            buf = BytesIO()
            pil_img.save(buf, format='PNG')
            file_bytes = buf.getvalue()
        except Exception:
            pass  # HEIC가 아니면 무시하고 진행

        # ===== [2] PIL 기반 이미지 디코딩 안정화 =====
        # PIL로 먼저 열어서 RGB 모드로 강제 변환 → EXIF 회전 보정 + 메타데이터 충돌 방지
        try:
            from PIL import ImageOps
            pil_img = Image.open(BytesIO(file_bytes))
            pil_img = ImageOps.exif_transpose(pil_img)  # EXIF 회전 보정
            pil_img = pil_img.convert('RGB')  # RGB 강제 변환 (RGBA, P, L 등 모든 모드 대응)
            buf = BytesIO()
            pil_img.save(buf, format='PNG')
            file_bytes = buf.getvalue()
        except Exception:
            pass

        # ===== [3] 바이트 → OpenCV =====
        nparr = np.frombuffer(file_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img is None:
            return file_bytes, 'png'

        # 4) Grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # 5) CLAHE 대비 향상 (그림자/조명 보정)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)

        # 6) Hough Line Transform 기반 deskew
        edges = cv2.Canny(enhanced, 50, 150, apertureSize=3)
        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=100,
                                minLineLength=img.shape[1] // 4, maxLineGap=10)
        if lines is not None and len(lines) > 5:
            angles = []
            for line in lines:
                x1, y1, x2, y2 = line[0]
                angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
                if abs(angle) < 10:  # 거의 수평인 선만
                    angles.append(angle)
            if angles:
                median_angle = np.median(angles)
                if abs(median_angle) > 0.3 and abs(median_angle) < 5:
                    (h, w) = enhanced.shape
                    center = (w // 2, h // 2)
                    M = cv2.getRotationMatrix2D(center, median_angle, 1.0)
                    enhanced = cv2.warpAffine(
                        enhanced, M, (w, h),
                        flags=cv2.INTER_CUBIC,
                        borderMode=cv2.BORDER_REPLICATE
                    )

        # 7) 약한 선명화 (과도한 sharpening 방지)
        blurred = cv2.GaussianBlur(enhanced, (0, 0), 2)
        sharpened = cv2.addWeighted(enhanced, 1.3, blurred, -0.3, 0)

        # PNG 인코딩
        _, encoded = cv2.imencode('.png', sharpened)
        return encoded.tobytes(), 'png'
    except Exception:
        return file_bytes, 'png'


# =====================================================================
# 📝 Fuzzy Matching 오타 보정 (경매 용어 특화)
# =====================================================================
# 경매 핵심 용어 사전 (OCR 오타 → 정확한 용어)
AUCTION_TERM_CORRECTIONS = {
    '가입류': '가압류', '가압유': '가압류', '가앞류': '가압류', '가압르': '가압류',
    '근저당관': '근저당권', '근저당귄': '근저당권', '근저당건': '근저당권',
    '근저당궈설정': '근저당권설정', '근저당관설정': '근저당권설정',
    '소유관이전': '소유권이전', '소유귄이전': '소유권이전', '소유건이전': '소유권이전',
    '채권최교액': '채권최고액', '채권쵀고액': '채권최고액', '채권최고엑': '채권최고액',
    '임의경매개시겸정': '임의경매개시결정', '임의경매겨시결정': '임의경매개시결정',
    '강제경매개시겸정': '강제경매개시결정',
    '전세관': '전세권', '전세귄': '전세권', '전세건': '전세권',
    '지상관': '지상권', '지상귄': '지상권', '지상건': '지상권',
    '유치관': '유치권', '유치귄': '유치권',
    '저당관': '저당권', '저당귄': '저당권',
    '담보가등거': '담보가등기', '담보가등끼': '담보가등기',
    '소유관': '소유권', '소유귄': '소유권',
    '말쇄': '말소', '말세': '말소',
    '접수': '접수', '첩수': '접수',
    '순위벤호': '순위번호', '순위번헌': '순위번호',
    '등기뫽적': '등기목적', '등기목젹': '등기목적',
    '배당요규': '배당요구',
    '가처붐': '가처분', '가처뷴': '가처분',
    '압류': '압류', '압유': '압류',
    '경매개시겸정': '경매개시결정',
}

@lru_cache(maxsize=None)
def _levenshtein_distance(s1, s2):
    """두 문자열 사이의 Levenshtein 편집 거리를 계산합니다."""
    if len(s1) < len(s2):
        return _levenshtein_distance(s2, s1)
    if len(s2) == 0:
        return len(s1)
    prev_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        curr_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = prev_row[j + 1] + 1
            deletions = curr_row[j] + 1
            substitutions = prev_row[j] + (c1 != c2)
            curr_row.append(min(insertions, deletions, substitutions))
        prev_row = curr_row
    return prev_row[-1]

def fuzzy_clean_text(text):
    """OCR 인식 결과에서 경매 용어 오타를 유사도 기반으로 자동 보정합니다."""
    # 1단계: 정확한 매칭 (빠른 경로)
    for wrong, correct in AUCTION_TERM_CORRECTIONS.items():
        if wrong in text:
            text = text.replace(wrong, correct)

    # 2단계: Fuzzy 매칭 (편집거리 1~2 이내)
    # 텍스트를 공백 기준으로 분리하여 각 토큰 검사
    words = text.split()
    corrected_words = []
    # 정확한 용어 목록 (중복 제거)
    correct_terms = list(set(AUCTION_TERM_CORRECTIONS.values()))

    for word in words:
        if len(word) >= 3:  # ⚠️ 3글자 이상만 Fuzzy 검사 (2글자 이하 '지분','말소' 등은 오변환 방지)
            best_match = None
            best_dist = float('inf')
            for term in correct_terms:
                # 길이 차이가 2 이상이면 스킵 (효율성)
                if abs(len(word) - len(term)) > 2:
                    continue
                dist = _levenshtein_distance(word, term)
                if dist <= 2 and dist < best_dist and dist > 0:
                    best_dist = dist
                    best_match = term
            if best_match:
                corrected_words.append(best_match)
            else:
                corrected_words.append(word)
        else:
            corrected_words.append(word)

    return ' '.join(corrected_words)


def normalize_names(parsed_records):
    """동일인으로 추정되는 이름(편집거리 ≤1)을 가장 빈도수가 높은 이름으로 통일합니다.
    OCR 오독(예: 조충희/조충회)을 방지합니다."""
    # 1단계: 전체 레코드에서 한글 이름 후보 수집
    name_pattern = re.compile(r'([가-힣]{2,4})(?:\s|$)')
    # 제외할 단어들 (권리 용어 등)
    exclude_words = set(list(AUCTION_TERM_CORRECTIONS.values()) + [
        '갑구', '을구', '표제부', '소유권', '근저당권', '가압류', '압류', '전세권',
        '지상권', '유치권', '저당권', '가처분', '가등기', '임차권', '경매', '말소',
        '설정', '이전', '변경', '접수', '등기', '목적', '권리자', '채무자', '채권자',
        '대한민국', '서울특별시', '경기도', '인천광역시', '부산광역시',
    ])

    name_freq = {}  # {이름: 빈도}
    for rec in parsed_records:
        content = rec.get('전체내용', '')
        names = name_pattern.findall(content)
        for name in names:
            if name not in exclude_words and len(name) >= 2:
                name_freq[name] = name_freq.get(name, 0) + 1

    if not name_freq:
        return parsed_records

    # 2단계: 유사 이름 그룹화 (편집거리 ≤ 1)
    name_list = list(name_freq.keys())
    name_groups = []  # [{이름1, 이름2, ...}, ...]
    assigned = set()

    for i, name_a in enumerate(name_list):
        if name_a in assigned:
            continue
        group = {name_a}
        for j in range(i + 1, len(name_list)):
            name_b = name_list[j]
            if name_b in assigned:
                continue
            if len(name_a) == len(name_b) and _levenshtein_distance(name_a, name_b) <= 1:
                group.add(name_b)
        if len(group) > 1:
            name_groups.append(group)
            assigned.update(group)

    # 3단계: 각 그룹에서 빈도 최고 이름을 대표로 선정하고 치환
    replacement_map = {}  # {오타이름: 대표이름}
    for group in name_groups:
        representative = max(group, key=lambda n: name_freq.get(n, 0))
        for name in group:
            if name != representative:
                replacement_map[name] = representative

    if not replacement_map:
        return parsed_records

    # 4단계: 레코드 내용에서 치환 적용
    for rec in parsed_records:
        content = rec.get('전체내용', '')
        for wrong_name, correct_name in replacement_map.items():
            content = content.replace(wrong_name, correct_name)
        rec['전체내용'] = content

    return parsed_records

# =====================================================================
# 🔑 2. API 키 자동 로드 (스트림릿 비밀 금고)
# =====================================================================
try:
    NAVER_API_URL = st.secrets["NAVER_API_URL"]
    NAVER_SECRET_KEY = st.secrets["NAVER_SECRET_KEY"]
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except KeyError:
    st.error("서버 관리자 설정 필요: 비밀 금고(Secrets)에 API 키가 설정되지 않았습니다.")
    st.stop()

# =====================================================================
# 🧠 3. 지식 베이스 및 필터링 룰 (개발자님 작성 원본 100% 반영)
# =====================================================================
base_keywords = ['근저당', '저당', '담보물권', '가압류', '압류', '체납처분압류', '강제경매개시결정', '임의경매개시결정', '경매개시결정', '담보가등기']
always_keep_keywords = ['건물철거', '토지인도', '법정지상권', '관습법상', '관습상', '분묘기지권', '예고등기', '요역지', '지역권', '도시철도법', '구분지상권', '채무자회생법', '특별매각조건', '인수조건']
ai_check_keywords = ['전세권', '임차권', '가처분', '처분금지가처분', '가등기', '소유권이전청구권', '지상권', '부합물', '종물', '다른 약정', '특약']

# 꼬리말 안내문 제거 필터
ignore_keywords = ["관할등기소", "본등기사항증명서", "사법부내부", "열람용이므로", "법적인효력", "실선으로그어진", "말소사항을표시", "기록사항없는", "기록사항없음", "열람일시", "사법부 말소사항", "수원지방법원", "고유번호", "말소사항을포함", "현재유효사항"]

# 📋 문서 유형 판별 키워드
registry_keywords = ["등기사항전부증명서", "등기사항증명서", "등기부등본", "갑구", "을구", "표제부"]
spec_keywords = ["매각물건명세서", "최선순위설정", "배당요구종기", "매각으로효력이", "임차인현황", "매각물건의표시"]

# 🚨 매각물건명세서 비고란 위험 키워드
danger_keywords = {
    "유치권": "🚨 [유치권 경고] 유치권 신고가 있습니다. 낙찰대금 외에 공사대금 등을 전액 떠안을 수 있으며, 변제할 때까지 건물 인도를 받지 못할 수 있습니다.",
    "법정지상권": "🚨 [법정지상권 경고] 토지/건물 소유권 분리로 인해 지료 분쟁이나 건물 철거 제약을 받을 수 있습니다.",
    "분묘기지권": "🚨 [분묘기지권 경고] 토지 위에 분묘가 있어 토지 사용에 심각한 제약이 있습니다.",
    "대지권미등기": "⚠️ [대지권 미등기 경고] 땅에 대한 권리가 불확실하여 대출 거절 및 추가 비용이 발생할 수 있습니다.",
    "토지별도등기": "⚠️ [토지별도등기 경고] 대지에 별도의 근저당 등이 설정되어 추가 부담이 있을 수 있습니다.",
    "농지취득자격증명": "❗ [보증금 몰수 경고] 농취증 미제출 시 입찰 보증금 전액 몰수됩니다. 입찰 전 발급 가능 여부를 반드시 확인하세요.",
    "위반건축물": "⚠️ [위반건축물 경고] 매년 이행강제금이 부과되며 원상복구 의무가 있을 수 있습니다.",
    "건물철거": "⛔ [건물 철거 위험] 토지 소유자의 건물 철거 소송으로 건물을 잃을 수 있습니다.",
    "대항력포기": "💡 [대항력 포기 확약서] HUG 등이 대항력 포기서를 제출하여 선순위 임차인 인수 부담이 없을 수 있습니다.",
    "원상회복": "⚠️ [원상회복 의무] 불법 형질변경 부분의 복구 비용을 낙찰자가 부담해야 합니다.",
    "제시외건물": "⚠️ [제시외 건물] 매각 대상에 포함되지 않는 건물이 있어 분쟁 위험이 있습니다.",
    "소유권상실": "⛔ [소유권 상실 위험] 선순위 가등기/가처분으로 소유권을 빼앗길 수 있습니다.",
}

knowledge_base = """
[경매 권리분석 핵심 규칙 요약 — AI 보조 판단용]

1. 말소기준권리 후보: 근저당권, 저당권, 압류, 가압류, 담보가등기, 경매개시결정등기, (배당요구한 최선순위 전세권)
2. 무조건 말소: 근저당권, 저당권, 압류, 가압류, 담보가등기 → 날짜 불문 무조건 소멸
3. 무조건 인수 (날짜 불문):
   - 유치권, 법정지상권, 관습상 법정지상권, 분묘기지권, 예고등기
   - 건물철거 및 토지인도청구권 보전 가처분 (후순위라도 절대 인수)
   - 도시철도법 등 공익사업 구분지상권
   - 채무자회생법 등기, 특별매각조건 인수 권리

4. 조건부 판단 (AI가 집중해야 할 권리):
   A. 전세권: 선순위 + 배당요구 → 소멸 (일부배당도 전액말소). 선순위 + 미배당요구 → 전액 인수.
      ※ 겸유 임차인(전세권자+임차인)이 임차인으로만 배당요구 → 전세권은 인수됨
   B. 임차권: 대항력발생일(전입신고 다음날) < 말소기준일 → 선순위.
      배당요구 안함 → 전액 인수. 배당요구+전액배당 → 소멸. 배당요구+일부배당 → 잔액 인수.
      ※ HUG 대항력 포기 확약서 제출 시 → 소멸 (Override)
   C. 가등기: 담보가등기 → 소멸. 소유권이전청구권보전가등기(선순위) → 인수 (소유권 상실 위험)
   D. 가처분: 선순위 → 인수. 후순위 → 소멸. 단, 건물철거/토지인도 → 무조건 인수.
   E. 지상권: 담보지상권(근저당과 함께 설정) → 소멸. 일반 선순위 → 인수.

5. 매각물건명세서 Override 규칙:
   - '대항력 포기' / '확약서' / '무상거주확인서' → 선순위 임차권도 소멸
   - '특별매각조건 인수' → 후순위라도 인수
   - '유치권 신고' → 무조건 인수 경고
   - '농지취득자격증명' → 보증금 몰수 위험 경고
"""

# =====================================================================
# 📋 문서 분류 및 명세서 분석 함수
# =====================================================================
def classify_document(ocr_text, use_title_only=False):
    """OCR 텍스트를 기반으로 등기부등본 vs 매각물건명세서 자동 판별.
    use_title_only=True: 문서 제목(상위 5행)만으로 판별 (더 정확)
    """
    clean = ocr_text.replace(" ", "")
    
    # 제목 영역 판별 (상위 5행 = 문서 제목 부분)
    if use_title_only:
        lines = ocr_text.strip().split("\n")
        title_area = " ".join(lines[:5]).replace(" ", "") if lines else clean
    else:
        title_area = clean
    
    # 제목 영역에서 강력 키워드 우선 검사
    title_spec_keywords = ["매각물건명세서", "매각물건의표시"]
    title_reg_keywords = ["등기사항전부증명서", "등기사항증명서", "등기부등본", "표제부"]
    
    for kw in title_spec_keywords:
        if kw in title_area:
            return "매각물건명세서"
    for kw in title_reg_keywords:
        if kw in title_area:
            return "등기부등본"
    
    # 전체 텍스트 스코어링 (fallback)
    spec_score = sum(1 for kw in spec_keywords if kw in clean)
    reg_score = sum(1 for kw in registry_keywords if kw in clean)
    # "갑구", "을구"는 명세서에도 등장할 수 있으므로 가중치 낮춤
    weak_reg_keywords = ["갑구", "을구"]
    strong_reg_count = sum(1 for kw in registry_keywords if kw not in weak_reg_keywords and kw in clean)
    
    if spec_score >= 2:
        return "매각물건명세서"
    if strong_reg_count >= 1:
        return "등기부등본"
    if reg_score >= 1:
        return "등기부등본"
    # 기본값: 등기부등본 (기존 호환성)
    return "등기부등본"

def detect_dangers(spec_text):
    """매각물건명세서 텍스트에서 위험 키워드를 탐지하여 경고 메시지 반환"""
    clean = spec_text.replace(" ", "")
    warnings = []
    for keyword, message in danger_keywords.items():
        if keyword in clean:
            warnings.append(message)
    return warnings

def ask_gemini_for_spec(spec_text, model):
    """매각물건명세서 OCR 텍스트에서 Gemini로 핵심 정보를 구조화 추출"""
    prompt = f"""
    너는 대한민국 법원 경매 전문가야.
    아래는 매각물건명세서를 OCR로 읽어온 텍스트야. 핵심 정보를 정확히 추출해 줘.

    [매각물건명세서 OCR 텍스트]
    {spec_text}

    [추출할 항목] - 해당 정보가 없으면 "확인불가"라고 적어줘.
    1. 최선순위 설정일자: (날짜와 권리명)
    2. 배당요구종기일: (날짜)
    3. 임차인 현황: (각 임차인별로 - 이름, 전입일, 확정일자, 보증금, 배당요구 여부)
    4. 매각으로 소멸되지 않는 권리: (목록)
    5. 비고란 특이사항: (전체 내용 요약)
    6. 특별매각조건: (있으면 기재)

    [출력 형식]
    각 항목을 번호와 함께 줄바꿈으로 구분해서 적어줘. 간결하게 핵심만 적어.
    """
    try:
        return model.generate_content(prompt).text
    except Exception:
        return "매각물건명세서 분석 실패"

def ask_gemini_for_rights(record_text, base_date, model, spec_summary=None, section_gu=None):
    spec_section = ""
    if spec_summary:
        spec_section = f"""

    [매각물건명세서 교차 검증 정보]
    {spec_summary}

    [교차 검증 추가 지시사항]
    - 매각물건명세서의 '최선순위 설정일'과 등기부의 말소기준권리 일자가 일치하는지 확인해.
    - 전세권자의 '배당요구 여부'가 명세서에 나와 있으면 확정 판단해 (배당요구 완료 → 말소).
    - 임차인의 전입신고일과 최선순위 설정일을 비교하여 대항력 유무를 판단해.
    - 명세서 정보가 있으므로 가능한 한 "추가확인" 대신 확정 판단(인수/말소)을 내려줘.
    - 이유 설명 시 "매각물건명세서상 ~로 확인되어"라는 근거를 포함해.
        """

    # 갑구/을구 구분 명시 (AI 혼동 원천 차단)
    section_note = ""
    if section_gu:
        section_note = f"\n    ⚠️ 중요: 이 데이터는 확실히 [{section_gu}] 섹션에서 추출된 데이터입니다. 갑구(소유권 관련)와 을구(소유권 이외의 권리)를 혼동하지 마세요."

    prompt = f"""
    너는 대한민국 법원 경매 권리분석 최고 전문가야.
    아래 [권리분석 예외 규칙]을 완벽하게 숙지하고, 제공된 [등기 권리 내용]을 분석해 줘.
    너의 외부 지식에 의존하지 말고, 오직 내가 제공한 [권리분석 예외 규칙]에 입각해서만 판단해.

    [권리분석 예외 규칙]
    {knowledge_base}

    [사건 기준 정보]
    - 확정된 말소기준권 일자: {base_date}
{section_note}
    [분석할 등기 권리 내용]
    {record_text}
    {spec_section}
    [지시사항]
    1. 위 등기 권리가 경매 낙찰 시 매수인에게 인수되는지, 아니면 말소되는지 판단해.
    2. 만약 등기부등본 내용만으로 확정할 수 없는 항목이라면 "추가확인 필요"라고 답변해.
    3. 출력 형식은 반드시 첫 줄에 "결과: 인수", "결과: 말소", "결과: 추가확인" 중 하나만 적고, 두 번째 줄에 "이유: (간략한 1~2줄 설명)"을 적어줘.
    """
    return model.generate_content(prompt).text

# 📐 Gemini Structured Output 스키마 정의 (환각 원천 차단)
class RightResult(typing.TypedDict):
    index: int
    result: str
    reason: str

def ask_gemini_for_rights_batch(ai_rows_data, base_date, model, spec_summary=None):
    """여러 건의 AI 정밀해석 대상을 단 1회의 API 호출로 일괄 처리합니다.
    ai_rows_data: list of dict — [{"index": int, "content": str, "section_gu": str}, ...]
    반환: dict — {index: {"결과": str, "이유": str}, ...}
    Gemini Structured Outputs를 사용하여 JSON 스키마를 강제 적용합니다.
    """
    if not ai_rows_data:
        return {}

    spec_section = ""
    if spec_summary:
        spec_section = f"\n[매각물건명세서 교차 검증 정보]\n{spec_summary}\n"

    # 사용자 특별 요청사항 반영
    user_req_section = ""
    if hasattr(st, 'session_state') and st.session_state.get('user_requests'):
        user_req_text = "\n".join([f"  {i+1}. {r}" for i, r in enumerate(st.session_state.user_requests)])
        user_req_section = f"\n[사용자 특별 요청사항 — 반드시 분석에 반영할 것]\n{user_req_text}\n"

    # 각 건을 번호로 태그하여 하나의 텍스트로 합침
    items_text = ""
    for item in ai_rows_data:
        section_note = f" (구분: {item['section_gu']})" if item.get('section_gu') else ""
        items_text += f"\n--- 항목 #{item['index']} {section_note} ---\n{item['content']}\n"

    prompt = f"""
    너는 대한민국 법원 경매 권리분석 최고 전문가야.
    아래 [권리분석 규칙]을 숙지하고, 여러 건의 등기 권리를 한꺼번에 분석해 줘.

    [권리분석 규칙]
    {knowledge_base}

    [사건 기준 정보]
    - 확정된 말소기준권리 일자: {base_date}
    {spec_section}
    {user_req_section}
    [분석할 등기 권리 목록 (총 {len(ai_rows_data)}건)]
    {items_text}

    [지시사항]
    1. 각 항목(#번호)에 대해 인수/말소/추가확인을 판단해.
    2. index에는 항목번호, result에는 "인수" 또는 "말소" 또는 "추가확인", reason에는 간략한 1~2줄 이유를 적어.
    """

    try:
        # Gemini Structured Outputs: response_schema로 JSON 형식 강제
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                response_mime_type="application/json",
                response_schema=list[RightResult]
            )
        )

        results_list = json.loads(response.text)

        results = {}
        for item in results_list:
            idx = item.get('index')
            result_str = item.get('result', '추가확인')
            reason = item.get('reason', '')
            if idx is not None:
                results[idx] = {'결과': result_str, '이유': reason}
        return results

    except (json.JSONDecodeError, Exception) as e:
        # JSON 파싱 실패 시 모든 항목을 수동 확인으로 처리
        fallback = {}
        for item in ai_rows_data:
            fallback[item['index']] = {'결과': '추가확인', '이유': f'⚠️ 판단 지연(수동 확인) — AI 응답 파싱 실패: {str(e)[:50]}'}
        return fallback

# =====================================================================
# 🔬 고급 교차 검증 로직 (등기부 전체 + 매각물건명세서 교차 분석)
# =====================================================================

def detect_daewi_risk(df, parsed_records):
    """대위변제 위험 감지: 1순위 근저당 금액이 소액이고
    후순위에 임차인 보증금이 클 때 경고를 생성합니다.
    대위변제 = 후순위 임차인이 선순위 근저당 채무를 대신 갚고
    근저당을 이전받아 경매를 취소시킬 수 있는 위험.
    """
    warnings = []
    try:
        all_text = ' '.join([str(r.get('전체내용', '')) for r in parsed_records])

        # 1순위 근저당(말소기준권리) 채권최고액 추출
        mortgage_amounts = []
        for _, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            purpose = str(row.get('등기목적', ''))
            if row.get('말소후보', False) and ('근저당' in purpose or '근저당' in content):
                amt_match = re.search(r'(?:채권최고액|금)\s*([\d,]+)\s*원', content.replace(' ', ''))
                if amt_match:
                    try:
                        mortgage_amounts.append(int(amt_match.group(1).replace(',', '')))
                    except ValueError:
                        pass

        # 임차인/전세권 보증금 추출
        tenant_amounts = []
        for _, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            purpose = str(row.get('등기목적', ''))
            if '임차' in purpose or '전세' in purpose or '임차' in content or '전세' in content:
                amt_match = re.search(r'(?:보증금|전세금|금)\s*([\d,]+)\s*원', content.replace(' ', ''))
                if amt_match:
                    try:
                        tenant_amounts.append(int(amt_match.group(1).replace(',', '')))
                    except ValueError:
                        pass

        if mortgage_amounts and tenant_amounts:
            min_mortgage = min(mortgage_amounts)
            max_tenant = max(tenant_amounts)
            # 1순위 근저당이 임차인 보증금의 30% 미만이면 대위변제 위험
            if min_mortgage > 0 and max_tenant > min_mortgage * 3:
                warnings.append(
                    f"⚠️ [대위변제 위험] 1순위 근저당 채권최고액({min_mortgage:,}원)이 "
                    f"후순위 임차인 보증금({max_tenant:,}원)보다 현저히 작습니다. "
                    f"임차인이 근저당 채무를 대위변제하여 경매를 취소시킬 수 있는 위험이 있습니다."
                )
    except Exception:
        pass
    return warnings

def detect_tax_seizure_conflict(df, parsed_records):
    """조세채권(당해세) 충돌 감지: '압류' 중 체납처분(세금) 압류가 있으면
    당해세가 임차인보다 우선 배당되어 임차인 미배당 잔액(인수액)이 증가할 위험을 경고.
    """
    warnings = []
    try:
        has_tax_seizure = False
        has_senior_tenant = False

        for _, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            purpose = str(row.get('등기목적', ''))
            result = str(row.get('결과', ''))

            # 체납처분 압류 (세금 관련) 감지
            if ('압류' in purpose or '압류' in content) and any(k in content for k in [
                '체납처분', '국세', '지방세', '세무서', '시청', '구청', '세금', '당해세',
                '재산세', '종합부동산세', '취득세', '양도소득세'
            ]):
                has_tax_seizure = True

            # 선순위 임차인 (인수 판정된) 감지
            if ('임차' in content or '전세' in content) and '인수' in result:
                has_senior_tenant = True

        if has_tax_seizure and has_senior_tenant:
            warnings.append(
                "⚠️ [당해세 충돌 경고] 체납처분 압류(세금)가 발견되었습니다. "
                "당해세(재산세·종부세 등)는 법정기일에 관계없이 근저당권보다 우선 배당되므로, "
                "선순위 임차인에게 돌아갈 배당금이 줄어들어 매수인의 인수액이 증가할 수 있습니다."
            )
        elif has_tax_seizure:
            warnings.append(
                "⚠️ [조세채권 경고] 체납처분 압류(세금)가 발견되었습니다. "
                "당해세는 우선 배당 대상이므로 배당 순위에 영향을 줄 수 있습니다."
            )
    except Exception:
        pass
    return warnings

def apply_spec_overrides(df, spec_summary):
    """매각물건명세서 Override: 비고란 핵심 키워드가 발견되면
    기존 determine_status 판단을 덮어씁니다.
    """
    warnings = []
    if not spec_summary:
        return warnings

    spec_text = str(spec_summary)

    try:
        # 1. 대항력 포기 / HUG 확약서 → 선순위 임차권/전세권 Override: 인수 → 말소
        if any(k in spec_text for k in ['대항력 포기', '대항력포기', '확약서', '무상거주확인서', '인수조건변경']):
            for idx, row in df.iterrows():
                content = str(row.get('전체내용', ''))
                result = str(row.get('결과', ''))
                if ('임차' in content or '전세' in content) and '인수' in result:
                    df.at[idx, '결과'] = "❌ 말소 (Override)"
                    df.at[idx, 'AI_상세이유'] = "매각물건명세서에 대항력 포기/확약서 기재 → 인수 → 말소로 변경"
            warnings.append("💡 [Override] 대항력 포기/확약서 확인됨 — 선순위 임차권도 소멸 처리")

        # 2. 토지별도등기 인수 → 후순위라도 강제 인수
        if any(k in spec_text for k in ['토지별도등기', '토지 별도등기', '별도등기 있음']):
            warnings.append(
                "⚠️ [토지별도등기 경고] 대지에 별도의 근저당 등이 설정되어 있어 "
                "건물 외에 토지 부분의 채무를 추가 인수할 위험이 있습니다."
            )

        # 3. 대지권 미등기
        if any(k in spec_text for k in ['대지권 미등기', '대지권미등기', '대지사용권 없음']):
            warnings.append(
                "⚠️ [대지권 미등기 경고] 집합건물의 대지권이 미등기 상태입니다. "
                "추가 분양대금 납부 또는 대지권 취득을 위한 추가 비용이 발생할 수 있습니다."
            )

        # 4. 특별매각조건 인수 → 말소 대상도 강제 인수 Override
        if any(k in spec_text for k in ['특별매각조건', '매수인이 인수', '인수하는 조건', '가압류등기의 부담을 매수인이 인수']):
            for idx, row in df.iterrows():
                content = str(row.get('전체내용', ''))
                result = str(row.get('결과', ''))
                if '말소' in result and any(k in content for k in ['가압류', '근저당', '압류']):
                    # 특별매각조건에 명시적으로 인수 대상으로 지정된 경우
                    if any(k in spec_text for k in [row.get('순위번호', '___NOMATCH___')]):
                        df.at[idx, '결과'] = "🚨 절대 인수 (Override)"
                        df.at[idx, 'AI_상세이유'] = "특별매각조건에 따라 말소 → 인수로 변경"
            warnings.append("⚠️ [Override] 특별매각조건이 발견되었습니다 — 일부 말소 대상이 인수로 변경될 수 있습니다.")

        # 5. 유치권 신고
        if any(k in spec_text for k in ['유치권 신고', '유치권 행사', '유치권 성립여부']):
            warnings.append(
                "🚨 [유치권 경고] 매각물건명세서에 유치권 신고가 기재되어 있습니다. "
                "낙찰대금 외에 공사대금 등을 전액 떠안을 수 있으며 건물 인도가 불가할 수 있습니다."
            )

        # 6. 농지취득자격증명
        if any(k in spec_text for k in ['농지취득자격증명', '농취증']):
            warnings.append(
                "❗ [보증금 몰수 경고] 농지취득자격증명 제출이 필요합니다. "
                "매각결정기일까지 미제출 시 입찰 보증금이 전액 몰수됩니다."
            )

        # 7. 법정지상권 성립 여지
        if any(k in spec_text for k in ['법정지상권 성립', '법정지상권이 성립할 여지']):
            warnings.append(
                "🚨 [법정지상권 경고] 매각물건명세서에 법정지상권 성립 여지가 기재되어 있습니다. "
                "토지 사용에 제한을 받을 수 있습니다."
            )

        # 8. 위반건축물 / 원상회복
        if any(k in spec_text for k in ['위반건축물', '원상회복', '불법형질변경']):
            warnings.append(
                "⚠️ [원상회복 경고] 위반건축물 또는 불법 형질변경이 확인되었습니다. "
                "이행강제금 부과 및 원상복구 비용을 매수인이 부담해야 합니다."
            )

        # 9. 건물만 매각 / 제시외 건물
        if any(k in spec_text for k in ['건물만 매각', '제시외 건물', '제시외건물']):
            warnings.append(
                "⚠️ [제시외 건물/건물만 매각] 대지 소유권이 없는 건물이거나 "
                "매각 대상에 포함되지 않는 제시외 건물이 있어 분쟁 위험이 있습니다."
            )

    except Exception:
        pass
    return warnings

def detect_share_auction_and_trust(df, parsed_records, all_clean_rows=None):
    """지분경매 및 신탁등기 감지:
    - 공유지분 경매 시 '공유자 우선매수청구권' 경고
    - 신탁등기 발견 시 '신탁원부 확인' 경고
    """
    warnings = []
    try:
        all_text = ' '.join([str(r.get('전체내용', '')) for r in parsed_records])
        # OCR 원본도 함께 검토
        if all_clean_rows:
            all_text += ' ' + ' '.join(all_clean_rows[:200])

        # 지분경매 감지
        share_keywords = ['지분', '공유', '분의', '2분의1', '3분의1', '4분의1',
                          '1/2', '1/3', '1/4', '공유자', '지분이전', '지분매각']
        if any(k in all_text for k in share_keywords):
            warnings.append(
                "⚠️ [지분경매 경고] 공유지분 관련 등기가 발견되었습니다. "
                "다른 공유자에게 '우선매수청구권'이 있어, 낙찰을 받아도 공유자가 같은 가격에 "
                "가져갈 수 있습니다. 또한 공유물분할청구 소송 위험도 있으니 신중하게 입찰하세요."
            )

        # 신탁등기 감지
        trust_keywords = ['신탁', '신탁원부', '신탁등기', '수탁자', '위탁자',
                          '자산관리공사', '우선수익자', '신탁계약']
        if any(k in all_text for k in trust_keywords):
            warnings.append(
                "⚠️ [신탁등기 경고] 신탁 관련 등기가 발견되었습니다. "
                "반드시 '신탁원부'를 열람하여 우선수익자, 신탁 조건, 처분 제한 여부를 확인하세요. "
                "신탁원부에 따라 매수인의 소유권 행사가 제한될 수 있습니다."
            )

    except Exception:
        pass
    return warnings

def detect_prev_owner_claims(df, parsed_records, spec_summary=None):
    """전 소유자 가압류/가처분 감지:
    채무자가 현재 소유자와 다른 가압류/가처분이 있으면
    매각물건명세서의 인수 조건을 교차 검증합니다.
    """
    warnings = []
    try:
        # 현재 소유자 이름 추출 (마지막 소유권이전 기록)
        current_owner = None
        for _, row in df.iterrows():
            purpose = str(row.get('등기목적', ''))
            if '소유권이전' in purpose or '소유권' in purpose:
                content = str(row.get('전체내용', ''))
                # 이름 추출 시도 (한글 2~4글자 이름 패턴)
                name_match = re.search(r'[가-힣]{2,4}(?:\s|$)', content)
                if name_match:
                    current_owner = name_match.group(0).strip()

        if current_owner:
            for idx, row in df.iterrows():
                purpose = str(row.get('등기목적', ''))
                content = str(row.get('전체내용', ''))
                result = str(row.get('결과', ''))

                if ('가압류' in purpose or '가처분' in purpose or
                    '가압류' in content or '가처분' in content):
                    # 채무자 이름 추출
                    debtor_match = re.search(r'채무자[:\s]*([\uac00-\ud7a3]{2,4})', content)
                    if debtor_match:
                        debtor_name = debtor_match.group(1).strip()
                        if debtor_name != current_owner:
                            # 전 소유자의 채무로 인한 가압류/가처분
                            spec_note = ""
                            if spec_summary and ('특별매각조건' in str(spec_summary) or '인수' in str(spec_summary)):
                                spec_note = " → 매각물건명세서에 특별매각조건/인수조건이 있어 인수될 수 있음"
                            warnings.append(
                                f"⚠️ [전 소유자 가압류/가처분] 현재 소유자({current_owner})와 "
                                f"채무자({debtor_name})가 다른 가압류/가처분이 발견되었습니다. "
                                f"전 소유자의 채무로 인한 권리로, 일반적으로는 말소되지만 "
                                f"특별매각조건에 따라 인수될 수 있으니 명세서를 확인하세요.{spec_note}"
                            )
                            break  # 첨 번째만 보고
    except Exception:
        pass
    return warnings

def detect_wage_claim_risk(df, parsed_records):
    """최우선 임금채권 경고: '임금', '근로복지공단' 등이 있으면서
    대항력 있는 임차인이 있을 때 인수액 폭탄 위험을 경고합니다.
    임금채권은 저당권보다 우선하여 임차인 배당금을 감소시탙니다.
    """
    warnings = []
    try:
        has_wage_claim = False
        has_senior_tenant = False
        wage_info = ""

        for _, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            purpose = str(row.get('등기목적', ''))
            result = str(row.get('결과', ''))

            # 임금채권 / 근로복지공단 감지
            if any(k in content for k in ['임금', '근로복지공단', '퇴직금', '임금채권',
                                          '임금등 체불금', '임금압류']):
                has_wage_claim = True
                wage_info = purpose

            # 대항력 있는 임차인 (인수 판정된) 감지
            if ('임차' in content or '전세' in content) and '인수' in result:
                has_senior_tenant = True

        if has_wage_claim and has_senior_tenant:
            warnings.append(
                f"🚨 [최우선 임금채권 경고] 임금/퇴직금 관련 채권({wage_info})이 발견되었습니다. "
                f"임금채권은 저당권보다 최우선 배당되므로, 선순위 임차인에게 돌아갈 배당금이 "
                f"대폭 감소하여 매수인의 인수액이 급격히 증가할 수 있습니다(인수액 폭탄 위험)."
            )
        elif has_wage_claim:
            warnings.append(
                f"⚠️ [임금채권 발견] 임금/퇴직금 관련 채권({wage_info})이 있습니다. "
                f"최우선 변제권으로 배당 순위에 영향을 줄 수 있습니다."
            )
    except Exception:
        pass
    return warnings

def detect_share_mortgage_scope(df, parsed_records, all_clean_rows=None):
    """지분경매 근저당 범위 판별:
    지분경매일 때 근저당권이 '부동산 전체'에 설정된 것인지
    '해당 지분'에만 설정된 것인지 구분하여 무잉여 취소나 인수 위험을 경고.
    """
    warnings = []
    try:
        all_text = ' '.join([str(r.get('전체내용', '')) for r in parsed_records])
        if all_clean_rows:
            all_text += ' ' + ' '.join(all_clean_rows[:200])

        # 지분경매 여부 확인
        is_share_auction = any(k in all_text for k in [
            '지분', '공유', '1/2', '1/3', '1/4', '2분의1', '3분의1',
            '지분매각', '지분경매', '공유자'
        ])

        if not is_share_auction:
            return warnings

        # 근저당이 전체 부동산에 설정되었는지 확인
        has_whole_property_mortgage = False
        has_share_mortgage = False

        for _, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            purpose = str(row.get('등기목적', ''))

            if '근저당' in purpose or '근저당' in content:
                # 전체 부동산 설정 키워드
                if any(k in content for k in ['부동산 전체', '전체부동산', '공동담보', '공동근저당']):
                    has_whole_property_mortgage = True
                # 지분만 설정 키워드
                if any(k in content for k in ['지분에 대하여', '지분에 관하여', '해당 지분',
                                              '지분근저당', '지분 근저당']):
                    has_share_mortgage = True

        if has_whole_property_mortgage:
            warnings.append(
                "🚨 [지분경매 근저당 전체설정] 근저당이 부동산 전체에 설정된 상태에서 "
                "지분만 경매됩니다. 매각대금이 근저당아액의 지분 분배액에 미치지 못하면 "
                "무잉여로 경매가 취소될 수 있으며, 근저당의 나머지 채무가 남을 수 있습니다."
            )
        elif has_share_mortgage:
            warnings.append(
                "⚠️ [지분경매 근저당 지분설정] 근저당이 해당 지분에만 설정되어 있습니다. "
                "전체 부동산 근저당보다 위험은 낮지만, 다른 공유자 지분에 별도 담보가 "
                "있는지 확인이 필요합니다."
            )
        elif is_share_auction:
            # 명시적 키워드가 없지만 지분경매임을 감지
            for _, row in df.iterrows():
                content = str(row.get('전체내용', ''))
                if '근저당' in content:
                    warnings.append(
                        "⚠️ [지분경매 근저당 범위 미상] 지분경매에서 근저당이 발견되었으나, "
                        "전체 부동산 설정인지 지분 설정인지 명확하지 않습니다. "
                        "등기부등본을 직접 확인하여 근저당 설정 범위를 파악하세요."
                    )
                    break

    except Exception:
        pass
    return warnings

def evaluate_eviction_difficulty(df, base_date, parsed_records):
    """명도 난이도 평가:
    대항력 없는 후순위 임차인이 '소액임차인 최우선변제' 대상이면
    명도가 수월하다는 코멘트를 추가합니다.
    """
    warnings = []
    try:
        for idx, row in df.iterrows():
            content = str(row.get('전체내용', ''))
            result = str(row.get('결과', ''))
            purpose = str(row.get('등기목적', ''))

            # 후순위 임차인 (말소 판정된 임차권)
            if ('임차' in content or '전세' in content) and '말소' in result:
                # 보증금액 추출
                amt_match = re.search(r'(?:보증금|전세금|금)\s*([\d,]+)\s*원', content.replace(' ', ''))
                if amt_match:
                    try:
                        deposit = int(amt_match.group(1).replace(',', ''))
                        # 소액임차인 최우선변제 기준 (2024년 기준 약 5,500만원 이하 — 지역별 다름)
                        # 대략적인 기준으로 1억 이하를 소액으로 간주
                        if deposit <= 100_000_000:
                            existing_reason = str(df.at[idx, 'AI_상세이유'])
                            eviction_note = (
                                f" 🟢 [명도 수월] 보증금 {deposit:,}원은 소액임차인 "
                                f"최우선변제 대상일 가능성이 높습니다. "
                                f"배당으로 보증금을 회수할 수 있어 자발적 퇴거 가능성이 높고 명도가 수월합니다."
                            )
                            df.at[idx, 'AI_상세이유'] = existing_reason + eviction_note
                        elif deposit > 100_000_000:
                            existing_reason = str(df.at[idx, 'AI_상세이유'])
                            eviction_note = (
                                f" 🟡 [명도 주의] 보증금 {deposit:,}원은 소액임차인 "
                                f"최우선변제 대상이 아닐 수 있어 명도소송이 필요할 수 있습니다."
                            )
                            df.at[idx, 'AI_상세이유'] = existing_reason + eviction_note
                    except ValueError:
                        pass
    except Exception:
        pass
    return warnings

def ask_gemini_for_malso_omission(all_records_text, base_date, model, spec_summary=None, confirmed_malso_summary=None):
    """Gemini에게 '매각 후 소멸(말소) 대상 권리 분석' 미션을 부여하여,
    매각 시 소멸되어야 할 권리와 인수 주의 권리를 분류합니다."""
    spec_ref = ""
    if spec_summary:
        spec_ref = f"""
    [매각물건명세서 참고 정보]
    {spec_summary}
        """

    # 확정된 말소 목록이 있으면 참조 데이터로 제공
    confirmed_ref = ""
    if confirmed_malso_summary:
        confirmed_ref = f"""
    [✅ 파이썬 엔진이 이미 확정한 말소 대상 등기 목록 (검증 완료)]
    아래는 프로그램이 하드코딩된 법률 규칙으로 이미 확정한 말소 대상입니다. 이 목록을 기준으로 삼아서 분석해 줘.
    {confirmed_malso_summary}
        """



    prompt = f"""
    너는 대한민국 법원 경매 권리분석 최고 전문가이자 '매각 후 소멸 대상 권리 분석 전문가'야.
    아래 등기부등본 전체 내용과 말소기준권리 일자를 바탕으로, **특수 미션**을 수행해 줘.

    [🎯 특수 미션: 매각 후 소멸(말소) 대상 권리 분석]
    경매 매각 시 말소촉탁으로 **소멸되어야 할 권리**를 분석하고, 현재 등기부에 여전히 살아있는 것처럼 보이는 권리가 있다면 알려줘.

    [권리분석 핵심 규칙]
    - 말소기준권리 후보: 근저당권, 저당권, 압류, 가압류, 담보가등기, 경매개시결정등기
    - 무조건 말소: 근저당권, 저당권, 압류, 가압류, 담보가등기 → 날짜 불문 소멸
    - 무조건 인수 (날짜 불문): 유치권, 법정지상권, 분묘기지권, 예고등기, 건물철거/토지인도 가처분, 채무자회생법 등기
    - 말소기준권리 이후 접수된 후순위 권리: 매각으로 소멸

    [사건 기준 정보]
    - 확정된 말소기준권리 일자: {base_date}
    {confirmed_ref}
    [등기부등본 전체 내용]
    {all_records_text}
    {spec_ref}
    [분석 기준]
    1. 말소기준권리 일자 이후에 접수된 등기 중, 매각 시 소멸 대상인데 OCR 텍스트에 '말소' 표시가 없거나 가로줄이 감지되지 않아 살아있는 것처럼 보이는 권리를 찾아줘.
    2. 단, '건물철거 및 토지인도청구권 보전 가처분', '예고등기', '유치권', '법정지상권', '분묘기지권', '도시철도법 구분지상권', '채무자회생법 등기', '특별매각조건 인수 권리' 등 예외적으로 소멸하지 않는 권리는 소멸 대상으로 판단하지 마. 이런 권리는 '인수 주의 권리'로 별도 표시해.
    3. 실제로 이미 말소된 등기(OCR 텍스트에 '말소' 단어가 포함되거나 취소선이 있는 경우)는 제외해.

    [출력 형식]
    - 매각 시 소멸 대상이 모두 정상 처리된 경우: "✅ 매각 후 소멸 대상 권리가 모두 정상적으로 확인되었습니다."
    - 소멸 대상인데 살아있는 것처럼 보이는 권리가 있으면 각 건에 대해:
      "🔶 매각 시 소멸 확정: [구분(갑구/을구)] 순위번호 [번호] - [등기목적] (접수일: [날짜])"
      "   사유: [왜 소멸되어야 하는지 간략 설명]"
    - 인수 주의 권리가 있으면:
      "⚠️ 인수 주의 권리: [구분(갑구/을구)] 순위번호 [번호] - [등기목적] (접수일: [날짜])"
      "   사유: [인수되는 이유 간략 설명]"
    """
    try:
        return model.generate_content(prompt).text
    except Exception as e:
        return f"매각 후 소멸 대상 권리 분석 실패 (API 오류: {str(e)[:150]})"

def ask_gemini_for_safety_report(df, base_date, model, spec_summary=None, parsed_records=None, confirmed_malso_summary=None):
    """전체 분석 결과를 바탕으로 Gemini에게 입찰 안전도 종합 의견을 요청합니다.
    데이터가 일부 누락되더라도 리포트가 실패하지 않도록 강화된 예외 처리를 적용합니다."""
    try:
        # 인수 권리 요약 생성 (개별 try-except로 부분 실패 허용)
        data_warnings = []  # 데이터 부족 경고 수집

        try:
            insu_rows = df[df['결과'].str.contains('인수', na=False) & ~df['결과'].str.contains('이미', na=False)]
        except Exception:
            insu_rows = pd.DataFrame()
            data_warnings.append("인수 권리 필터링 실패")

        try:
            malso_rows = df[df['결과'].str.contains('말소', na=False) & ~df['결과'].str.contains('이미', na=False)]
        except Exception:
            malso_rows = pd.DataFrame()
            data_warnings.append("말소 권리 필터링 실패")

        try:
            danger_rows = df[df['결과'].str.contains('절대 인수', na=False)]
        except Exception:
            danger_rows = pd.DataFrame()
            data_warnings.append("위험 권리 필터링 실패")

        try:
            warning_rows = df[df['결과'].str.contains('서류확인', na=False)]
        except Exception:
            warning_rows = pd.DataFrame()
            data_warnings.append("서류확인 항목 필터링 실패")

        try:
            insu_summary = "\n".join(
                [f"  - {r['구분']} 순위번호 {r['순위번호']}: {r['등기목적']}" for _, r in insu_rows.iterrows()]
            ) if not insu_rows.empty else "  없음"
        except Exception:
            insu_summary = "  데이터 추출 실패"
            data_warnings.append("인수 권리 요약 생성 실패")

        try:
            danger_summary = "\n".join(
                [f"  - {r['구분']} 순위번호 {r['순위번호']}: {r['등기목적']}" for _, r in danger_rows.iterrows()]
            ) if not danger_rows.empty else "  없음"
        except Exception:
            danger_summary = "  데이터 추출 실패"
            data_warnings.append("위험 권리 요약 생성 실패")

        # 채권액 추출 시도 (근저당권 금액)
        amount_info = ""
        try:
            if parsed_records:
                amounts = []
                for rec in parsed_records:
                    content = rec.get('전체내용', '')
                    amt_matches = re.findall(r'(?:채권최고액|금)\s*([\d,]+)\s*원', content.replace(' ', ''))
                    for amt in amt_matches:
                        try:
                            amounts.append(int(amt.replace(',', '')))
                        except ValueError:
                            pass
                if amounts:
                    total_amt = sum(amounts)
                    amount_info = f"\n  - 감지된 채권최고액 합계: 약 {total_amt:,}원 ({len(amounts)}건)"
        except Exception:
            data_warnings.append("채권최고액 추출 실패")

        spec_ref = ""
        if spec_summary:
            spec_ref = f"\n[매각물건명세서 분석 결과]\n{spec_summary}"

        # 확정된 말소 목록 참조
        confirmed_ref = ""
        if confirmed_malso_summary:
            confirmed_ref = f"\n[✅ 확정된 말소 대상 등기 목록 (프로그램 검증 완료)]\n{confirmed_malso_summary}"

        # 사용자 특별 요청사항 반영
        user_req_ref = ""
        if hasattr(st, 'session_state') and st.session_state.get('user_requests'):
            user_req_text = "\n".join([f"  {i+1}. {r}" for i, r in enumerate(st.session_state.user_requests)])
            user_req_ref = f"\n[사용자 특별 요청사항 — 반드시 평가에 반영할 것]\n{user_req_text}"

        # 데이터 부족 경고 문구
        data_warning_note = ""
        if data_warnings:
            data_warning_note = f"\n\n    ⚠️ 데이터 부족 알림: 일부 데이터 추출에 실패했습니다 ({', '.join(data_warnings)}). 아래 분석은 확보된 데이터 범위 내의 추정치입니다."

        prompt = f"""
    너는 대한민국 법원 경매 권리분석 최고 전문가이자 투자 리스크 평가 전문가야.
    아래 분석 결과를 바탕으로 "이 물건에 입찰해도 안전한지" 종합 의견을 줘.
{data_warning_note}
    [분석 기준 정보]
    - 말소기준권리 일자: {base_date}
    - 총 등기 건수: {len(df)}건
    - 말소 예정: {len(malso_rows)}건
    - 인수 예정: {len(insu_rows)}건
    - 절대 인수 (위험): {len(danger_rows)}건
    - 서류확인 필요: {len(warning_rows)}건{amount_info}

    [인수되는 권리 목록]
{insu_summary}

    [절대 인수 (위험 권리) 목록]
{danger_summary}
    {spec_ref}
    {confirmed_ref}
    {user_req_ref}
    [지시사항]
    1. 위험도 등급을 반드시 첫 줄에 표시해: "🟢 안전", "🟡 주의", "🔴 위험" 중 하나.
       - 🟢 안전: 인수되는 위험 권리 없음, 매각 후 소멸 대상 정상 처리
       - 🟡 주의: 인수 권리 있으나 금액이 크지 않거나, 서류확인 필요 건이 있음
       - 🔴 위험: 절대 인수 권리 존재, 유치권/건물철거/법정지상권 등 중대 위험
    2. 두 번째 줄부터 간결한 종합 의견 (3~5줄):
       - 인수되는 채권의 총 부담 추정액
       - 핵심 위험 요소 요약
       - 입찰 시 주의사항
    3. 마지막에 "💡 입찰팁: " 으로 시작하는 실용적 조언 1줄 추가.
    4. 데이터가 부족한 경우에도 확보된 정보 범위 내에서 가능한 분석을 반드시 제시해. 부족한 부분은 '데이터 부족으로 인한 추정치'임을 명시해.
    """
        try:
            result = model.generate_content(prompt).text
            if data_warnings:
                result = f"⚠️ [일부 데이터 부족으로 인한 추정치입니다]\n\n{result}"
            return result
        except Exception as inner_e:
            return f"종합 안전도 리포트 생성 실패 (API 오류: {str(inner_e)[:150]})"
    except Exception as e:
        return f"⚠️ 종합 안전도 리포트 생성 중 오류가 발생했습니다: {e}\n\n확보된 데이터가 부족하여 리포트를 생성할 수 없습니다. 원본 등기부등본을 다시 확인해 주세요."



def merge_and_sort_pages(sorted_files, ocr_cache):
    """여러 장의 등기부등본 페이지를 갑구/을구 섹션 기반으로 자동 정렬합니다.
    페이지 순서가 섞여 있어도 표제부→갑구→을구 순서로 올바르게 재배열합니다."""
    section_keywords = {
        '표제부': ['표제부', '건물의표시', '토지의표시', '대지권의목적', '1동의건물의표시'],
        '갑구': ['갑구', '소유권에관한사항'],
        '을구': ['을구', '소유권이외의권리']
    }

    pages_info = []
    for file in sorted_files:
        raw_bytes = file.getvalue()
        file_hash = hashlib.sha256(raw_bytes).hexdigest()
        if file_hash not in ocr_cache:
            continue
        page_rows = ocr_cache[file_hash]
        page_text = ' '.join(page_rows).replace(' ', '')

        # 섹션 감지
        detected_section = None
        section_priority = {'표제부': 0, '갑구': 1, '을구': 2}
        for section, keywords in section_keywords.items():
            for kw in keywords:
                if kw in page_text:
                    if detected_section is None or section_priority[section] < section_priority.get(detected_section, 99):
                        detected_section = section
                    break

        # 순위번호 추출 (연속성 판단용)
        rank_numbers = re.findall(r'(?:^|\s)([1-9]\d{0,2})\s', ' '.join(page_rows))
        first_rank = int(rank_numbers[0]) if rank_numbers else 999

        pages_info.append({
            'file': file,
            'hash': file_hash,
            'section': detected_section or '갑구',  # 기본값: 갑구
            'first_rank': first_rank,
            'rows': page_rows
        })

    # 정렬: 표제부 → 갑구 → 을구, 같은 섹션 내에서는 순위번호 순
    section_order = {'표제부': 0, '갑구': 1, '을구': 2}
    pages_info.sort(key=lambda p: (section_order.get(p['section'], 1), p['first_rank']))

    # 정렬된 rows 반환
    merged_rows = []
    for page in pages_info:
        merged_rows.extend(page['rows'])

    return merged_rows

# =====================================================================
# 🔄 4. 화면 자동 전환 로직 (세션 상태 관리)
# =====================================================================
if 'step' not in st.session_state:
    st.session_state.step = 1  
if 'final_df' not in st.session_state:
    st.session_state.final_df = None
if 'malso_df' not in st.session_state:
    st.session_state.malso_df = None
if 'ocr_cache' not in st.session_state:
    st.session_state.ocr_cache = {}  # 💰 OCR 결과 캐시 (key: 파일 해시)
if 'spec_summary' not in st.session_state:
    st.session_state.spec_summary = None  # 📋 매각물건명세서 요약
if 'malso_omission_report' not in st.session_state:
    st.session_state.malso_omission_report = None  # 🔍 말소 누락 탐지 보고서
if 'danger_warnings' not in st.session_state:
    st.session_state.danger_warnings = []  # 🚨 위험 경고 목록
if 'cross_warnings' not in st.session_state:
    st.session_state.cross_warnings = []  # 🔬 고급 교차 검증 경고 목록
if 'base_date_info' not in st.session_state:
    st.session_state.base_date_info = None  # 📅 말소기준권리 상세 정보
if 'safety_report' not in st.session_state:
    st.session_state.safety_report = None  # 🧾 종합 안전도 리포트
if 'user_requests' not in st.session_state:
    st.session_state.user_requests = []  # 📝 사용자 요청사항 목록
if 'user_request_count' not in st.session_state:
    st.session_state.user_request_count = 0  # 📝 요청 횟수 카운터 (최대 2회)
if 'qa_history' not in st.session_state:
    st.session_state.qa_history = []  # 💬 Q&A 대화 기록
if 'qa_count' not in st.session_state:
    st.session_state.qa_count = 0  # 💬 질문 횟수 카운터 (최대 2회)



# =====================================================================
# 📱 [1단계 화면] 메인 화면 및 사진 업로드
# =====================================================================
if st.session_state.step == 1:
    st.title("🧙‍♂️ AI 경매 권리분석 마법사")
    st.markdown("스마트폰으로 등기부등본과 매각물건명세서 사진을 찍어서 올리면, AI가 자동으로 권리를 분석해 줍니다.")
    
    # CSS로 영어 문구가 완벽히 숨겨진 업로드 창
    uploaded_files = st.file_uploader(" ", accept_multiple_files=True, type=['jpg', 'jpeg', 'png', 'heic', 'heif'], label_visibility="collapsed", key="photo_uploader")

    # =========================================================
    # 📝 AI에게 요청할 사항 입력 영역 (최대 2회) — 권리분석 시작 바로 위
    # =========================================================
    with st.expander("📝 AI에게 요청할 사항 (선택사항)", expanded=bool(st.session_state.user_requests)):
        st.caption("권리분석 시 AI가 특별히 확인하거나 반영해야 할 사항을 입력해 주세요. (최대 2회)")

        # 이미 등록된 요청사항 표시
        if st.session_state.user_requests:
            for i, req in enumerate(st.session_state.user_requests, 1):
                st.info(f"📌 요청 {i}: {req}")

        if st.session_state.user_request_count < 2:
            user_req_input = st.text_area(
                "요청사항을 입력하세요",
                placeholder="예: 전세권이 있는 경우 보증금 반환 가능성을 자세히 분석해 주세요. / 이 물건은 토지만 경매입니다. / 임차인이 실제 거주 중입니다.",
                key="user_request_input",
                height=80,
                label_visibility="collapsed"
            )
            remaining = 2 - st.session_state.user_request_count
            if st.button(f"📝 요청사항 등록 (남은 횟수: {remaining}회)", use_container_width=True):
                if user_req_input and user_req_input.strip():
                    st.session_state.user_requests.append(user_req_input.strip())
                    st.session_state.user_request_count += 1
                    st.rerun()
                else:
                    st.warning("요청사항을 입력해 주세요.")
        else:
            st.success("✅ 요청사항 2회를 모두 등록하셨습니다. 등록된 내용은 AI 분석에 반영됩니다.")

    if st.button("🚀 권리분석 시작", type="primary", use_container_width=True):
        if not uploaded_files:
            st.warning("사진을 먼저 업로드해주세요.")
        else:
            try:
                genai.configure(api_key=GEMINI_API_KEY)
                model = genai.GenerativeModel('gemini-3-flash-preview')
                
                # 📂 Natural Sort: 파일명 숫자 기준 정렬 (1 < 2 < 10)
                def natural_sort_key(f):
                    return [int(c) if c.isdigit() else c.lower()
                            for c in re.split(r'(\d+)', f.name)]

                # 📊 진행률 표시 OCR 스캔
                all_clean_rows = []
                cache_hit_count = 0
                total_files = len(uploaded_files)
                sorted_files = sorted(uploaded_files, key=natural_sort_key)
                progress_bar = st.progress(0, text='📸 분석 준비 중...')


                # =====================================================
                # 🚀 비동기 OCR 처리 (aiohttp + asyncio.Semaphore)
                # =====================================================
                # 1단계: 캐시 히트 파일 사전 분리 + 전처리
                ocr_tasks = []  # [(file_idx, file_name, file_hash, preprocessed_bytes, file_format)]
                for file_idx, file in enumerate(sorted_files):
                    raw_bytes = file.getvalue()
                    file_hash = hashlib.sha256(raw_bytes).hexdigest()

                    if file_hash in st.session_state.ocr_cache:
                        progress_bar.progress((file_idx + 1) / total_files, text=f'💰 캐시 사용 ({file_idx + 1}/{total_files}): {file.name}')
                        all_clean_rows.extend(st.session_state.ocr_cache[file_hash])
                        cache_hit_count += 1
                    else:
                        # 전처리는 동기로 수행 (CPU 작업)
                        preprocessed_bytes, file_format = smart_preprocess(raw_bytes)
                        ocr_tasks.append((file_idx, file.name, file_hash, preprocessed_bytes, file_format))

                # 2단계: 비동기 OCR API 호출
                if ocr_tasks:
                    async def _ocr_single_file(session, sem, file_name, preprocessed_bytes, file_format, api_url, secret_key):
                        """단일 파일 비동기 OCR 호출 (Semaphore 제어)"""
                        async with sem:
                            request_json = {
                                'images': [{'format': file_format, 'name': 'demo'}],
                                'requestId': str(uuid.uuid4()),
                                'version': 'V2',
                                'timestamp': int(round(time.time() * 1000))
                            }
                            headers = {'X-OCR-SECRET': secret_key}
                            file_mime = 'image/png'

                            # 파일명 확장자를 실제 포맷(png)에 맞춤 (확장자 불일치 → HTTP 400 방지)
                            safe_filename = os.path.splitext(file_name)[0] + '.' + file_format

                            for attempt in range(3):
                                form_data = aiohttp.FormData()
                                # ⚠️ message는 반드시 문자열로 전송 (bytes로 보내면 aiohttp가 파일 업로드로 처리하여 API 400 에러 발생)
                                form_data.add_field('message', json.dumps(request_json))
                                form_data.add_field('file', preprocessed_bytes,
                                                    filename=safe_filename,
                                                    content_type=file_mime)
                                try:
                                    async with session.post(api_url, headers=headers, data=form_data, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                                        if resp.status == 200:
                                            return {'status': 'ok', 'data': await resp.json(), 'file_name': file_name}
                                        elif resp.status == 429 and attempt < 2:
                                            await asyncio.sleep(2 ** attempt)
                                            continue
                                        else:
                                            error_body = await resp.text()
                                            return {'status': 'error', 'code': resp.status, 'body': error_body, 'file_name': file_name}
                                except asyncio.TimeoutError:
                                    if attempt < 2:
                                        await asyncio.sleep(2)
                                        continue
                                    else:
                                        return {'status': 'timeout', 'file_name': file_name}
                                except Exception as e:
                                    return {'status': 'exception', 'error': str(e), 'file_name': file_name}
                            return {'status': 'error', 'code': 'max_retries', 'file_name': file_name}

                    async def _run_all_ocr(tasks, api_url, secret_key):
                        """모든 OCR 작업을 비동기로 실행 (최대 3개 동시)"""
                        sem = asyncio.Semaphore(3)
                        async with aiohttp.ClientSession() as session:
                            coroutines = [
                                _ocr_single_file(session, sem, t[1], t[3], t[4], api_url, secret_key)
                                for t in tasks
                            ]
                            return await asyncio.gather(*coroutines)

                    progress_bar.progress(cache_hit_count / total_files if total_files > 0 else 0, text='🚀 비동기 OCR 스캔 중...')
                    ocr_results = asyncio.run(_run_all_ocr(ocr_tasks, NAVER_API_URL, NAVER_SECRET_KEY))

                    # 3단계: 결과 처리 (텍스트 파싱 + 캐싱)
                    for task_idx, (file_idx, file_name, file_hash, _, _) in enumerate(ocr_tasks):
                        result = ocr_results[task_idx]
                        progress_bar.progress((file_idx + 1) / total_files, text=f'📸 처리 중 ({file_idx + 1}/{total_files}): {file_name}')

                        if result['status'] == 'timeout':
                            st.error(f"⏱️ OCR 응답 시간 초과: {file_name}")
                            st.stop()
                        elif result['status'] == 'exception':
                            st.error(f"🌐 네트워크 오류: {result.get('error', '')}")
                            st.stop()
                        elif result['status'] == 'error':
                            error_detail = result.get('body', '')[:200]
                            st.error(f"OCR 스캔 실패 ({file_name}): HTTP {result.get('code', 'Unknown')}\n{error_detail}")
                            st.stop()

                        # 성공 처리
                        images_data = result['data'].get('images', [])
                        fields = images_data[0].get('fields', []) if images_data else []
                        if not fields:
                            st.warning(f"⚠️ {file_name}에서 텍스트가 감지되지 않았습니다.")
                            continue

                        current_row, last_y, page_rows = [], -1, []
                        low_conf_words = []
                        sorted_fields = sorted(fields, key=lambda x: x['boundingPoly']['vertices'][0]['y'])

                        for field in sorted_fields:
                            text = field['inferText']
                            confidence = field.get('inferConfidence', 1.0)
                            y_pos = field['boundingPoly']['vertices'][0]['y']
                            x_pos = field['boundingPoly']['vertices'][0]['x']
                            text = re.sub(r'(\d{6})\s*-\s*\d{7}', r'\1-*******', text)

                            # 📝 저신뢰도 단어 → Fuzzy Matching 우선 보정
                            if confidence < 0.5:
                                original_text = text
                                text = fuzzy_clean_text(text)
                                if text != original_text:
                                    low_conf_words.append(f"'{original_text}'→'{text}' (신뢰도:{confidence:.2f})")

                            if last_y == -1 or abs(y_pos - last_y) <= 35:
                                current_row.append({'x': x_pos, 'text': text})
                            else:
                                current_row.sort(key=lambda x: x['x'])
                                page_rows.append(" ".join([item['text'] for item in current_row]))
                                current_row = [{'x': x_pos, 'text': text}]
                            last_y = y_pos

                        if current_row:
                            current_row.sort(key=lambda x: x['x'])
                            page_rows.append(" ".join([item['text'] for item in current_row]))

                        # 📝 전체 행에 Fuzzy 오타 보정 적용
                        page_rows = [fuzzy_clean_text(row) for row in page_rows]

                        if low_conf_words:
                            st.toast(f"📝 {file_name}: {len(low_conf_words)}개 저신뢰도 단어 자동 보정")

                        st.session_state.ocr_cache[file_hash] = page_rows
                        all_clean_rows.extend(page_rows)

                progress_bar.progress(1.0, text='✅ 분석 완료!')

                if cache_hit_count > 0:
                    st.toast(f'💰 {cache_hit_count}개 파일은 이전 분석 결과를 재사용했습니다 (부분 재분석 · 비용 절감!)')


                # 📷 Feature C: 여러 장 등기부 페이지 연결 — 갑구/을구 자동 정렬
                original_clean_rows = list(all_clean_rows)  # 원본 보존 (fallback용)
                if total_files > 1:
                    merged_rows = merge_and_sort_pages(sorted_files, st.session_state.ocr_cache)
                    if merged_rows and len(merged_rows) == len(all_clean_rows):
                        all_clean_rows = merged_rows
                        st.toast('📷 여러 페이지가 갑구/을구 순서로 자동 정렬되었습니다!')

                # 📋 문서 분류: 등기부등본 vs 매각물건명세서 자동 분리
                registry_rows, spec_rows = [], []

                # 파일별이 아닌 전체 텍스트 기반으로 분류 (더 정확)
                full_text = " ".join(all_clean_rows)
                doc_type = classify_document(full_text)

                if doc_type == "매각물건명세서":
                    # 전체가 명세서인 경우 (등기부 없이 명세서만 올린 경우)
                    spec_rows = all_clean_rows
                    registry_rows = []
                else:
                    # 파일별로 분류 시도
                    temp_registry, temp_spec = [], []
                    # 각 파일의 OCR 텍스트를 다시 가져와서 분류
                    offset = 0
                    for file in sorted(uploaded_files, key=natural_sort_key):
                        raw_bytes = file.getvalue()
                        file_hash = hashlib.sha256(raw_bytes).hexdigest()
                        if file_hash in st.session_state.ocr_cache:
                            file_rows = st.session_state.ocr_cache[file_hash]
                        else:
                            # 캐시에 없으면 전체에서 추정 (fallback)
                            file_rows = all_clean_rows[offset:]
                        file_text = " ".join(file_rows)
                        file_type = classify_document(file_text, use_title_only=True)
                        if file_type == "매각물건명세서":
                            temp_spec.extend(file_rows)
                        else:
                            temp_registry.extend(file_rows)
                        offset += len(file_rows)

                    if temp_spec:
                        registry_rows = temp_registry
                        spec_rows = temp_spec
                    else:
                        registry_rows = all_clean_rows
                        spec_rows = []

                # 📋 매각물건명세서가 있으면 Gemini로 구조화 추출
                spec_summary = None
                danger_warnings_list = []
                if spec_rows:
                    with st.spinner('📋 매각물건명세서를 분석하고 있습니다...'):
                        spec_full_text = "\n".join(spec_rows)
                        spec_summary = ask_gemini_for_spec(spec_full_text, model)
                        danger_warnings_list = detect_dangers(spec_full_text)
                        st.toast('📋 매각물건명세서가 감지되었습니다. 교차 검증을 수행합니다!')

                st.session_state.spec_summary = spec_summary
                st.session_state.danger_warnings = danger_warnings_list

                # 등기부등본 파싱 (registry_rows 또는 all_clean_rows 사용)
                analysis_rows = registry_rows if registry_rows else all_clean_rows

                with st.spinner('파이썬 엔진이 권리를 분류하고 있습니다...'):
                    def parse_records_from_rows(rows):
                        """OCR 텍스트 행들에서 등기 레코드를 파싱합니다."""
                        _records, _current_record, _current_gu = [], {}, None
                        # 순위번호: 기본 패턴 + 느슨한 fallback
                        _rank_pattern = re.compile(r'^([1-9]\d{0,2}(?:-[1-9]\d{0,2})?)(?:\s+|번|(?=[가-힣]))')
                        # 날짜: 3단계 fallback
                        _date_patterns = [
                            re.compile(r'(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일'),  # 2024년 01월 31일
                            re.compile(r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})'),  # 2024.01.31 / 2024-01-31
                            re.compile(r'(20\d{2})(\d{2})(\d{2})'),  # 20240131 (8자리 연속 숫자)
                        ]
                        # 접수번호: 3단계 fallback (강화)
                        _receipt_patterns = [
                            re.compile(r'제\s*(\d+)\s*호'),  # 제XXXXX호 (정상)
                            re.compile(r'(?<!번호)(?<!년)(?<!월)(?<!일)(\d{5,6})(?!년|월|일|호|번)'),  # 5~6자리 숫자 (날짜/번호 제외)
                            re.compile(r'.*?(\d{5,7}).*?'),  # 최후 fallback: '제'/'호' 깨져도 5~7자리 숫자 추출
                        ]

                        for row in rows:
                            clean_row = row.replace(" ", "")
                            if any(kw in clean_row for kw in ignore_keywords):
                                continue
                            if "갑구" in clean_row and ("소유권" in clean_row or "관한사항" in clean_row): _current_gu = "갑구"; continue
                            if "을구" in clean_row and ("소유권" in clean_row or "관한사항" in clean_row or "이외의권리" in clean_row): _current_gu = "을구"; continue
                            if _current_gu is None or "순위번호" in row or "등기목적" in row or "접수" in row: continue

                            match = _rank_pattern.match(row)
                            is_new_record = False

                            if match:
                                rank_str = match.group(1)
                                rest_of_line = row[match.end():].strip()

                                if "-" in rank_str:
                                    parts = rank_str.split('-')
                                    if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                                        parent_rank = int(parts[0])
                                        if parent_rank <= 200:
                                            is_new_record = True
                                elif len(rank_str) >= 4 or int(rank_str) > 200:
                                    pass
                                elif rest_of_line.startswith(('호', '동', '층', '길', '번지', 'm', '㎡', '전', '년', '월', '일')):
                                    pass
                                else:
                                    is_new_record = True

                            if is_new_record:
                                if _current_record: _records.append(_current_record)
                                # 갑구/을구 섹션 정보를 전체내용에 태깅 (AI 혼동 방지)
                                tagged_content = f"[{_current_gu}] {row}"
                                _current_record = {'구분': _current_gu, '순위번호': rank_str, '전체내용': tagged_content}
                            else:
                                if _current_record: _current_record['전체내용'] += " " + row

                        if _current_record: _records.append(_current_record)

                        _parsed = []
                        for rec in _records:
                            content = rec['전체내용'].replace(" ", "")
                            # 날짜 추출: 3단계 fallback
                            date_match = None
                            for dp in _date_patterns:
                                date_match = dp.search(rec['전체내용'])
                                if date_match:
                                    break
                            rec['접수일자_기준'] = None

                            if date_match:
                                y, m, d = date_match.groups()
                                try:
                                    rec['접수일자_기준'] = datetime.date(int(y), int(m), int(d))
                                except ValueError:
                                    pass

                                # 접수번호 추출: 3단계 fallback (강화)
                                receipt_num = None
                                for rp in _receipt_patterns:
                                    rm = rp.search(rec['전체내용'])
                                    if rm:
                                        candidate = rm.group(1)
                                        # 최후 fallback 패턴의 경우 날짜 숫자와 겹치지 않도록 추가 검증
                                        if len(candidate) >= 5 and len(candidate) <= 7:
                                            receipt_num = candidate
                                            break
                                        elif rp != _receipt_patterns[-1]:  # 최후 fallback이 아니면 그대로 사용
                                            receipt_num = candidate
                                            break
                                rec['접수번호'] = receipt_num  # 접수번호 별도 저장 (호버용)
                                rec['접수일자_표시'] = f"{y}년 {m}월 {d}일" + (f" 제{receipt_num}호" if receipt_num else "")

                                raw_target = rec['전체내용'][:date_match.start()].replace(rec['순위번호'], '', 1).strip()
                                action_strip_pattern = r'^번\s*|(?:전부)?근저당권설정$|가압류$|임의경매개시결정$|강제경매개시결정$|압류$|경매개시결정$'
                                clean_target = re.sub(action_strip_pattern, '', raw_target).strip()

                                action = ""
                                if '임의경매개시결정' in content: action = "임의경매개시결정"
                                elif '강제경매개시결정' in content: action = "강제경매개시결정"
                                elif '가압류' in content: action = "가압류"
                                elif '근저당권설정' in content: action = "전부근저당권설정" if '전부근저당권설정' in content else "근저당권설정"
                                elif '압류' in content: action = "압류"
                                if action and action in clean_target:
                                    rec['등기목적'] = clean_target
                                else:
                                    rec['등기목적'] = f"{clean_target} {action}".strip()
                            else:
                                rec['접수일자_표시'], rec['등기목적'] = "확인불가", "확인불가"

                            rec['말소후보'] = any(kw in content for kw in base_keywords)

                            # 🛡️ 교차 검증: 등기목적(purpose) 자체에 말소기준권리 키워드가 없으면 말소후보 해제
                            # (전체내용에 다른 등기의 키워드가 섞여 들어온 경우 방지)
                            # 예: 소유권이전, 파산선고, 상속, 증여 등은 전체내용에 '압류' 등이 포함되더라도 말소후보가 아님
                            if rec['말소후보']:
                                purpose_clean = rec.get('등기목적', '').replace(' ', '')
                                purpose_has_base_kw = any(kw in purpose_clean for kw in base_keywords)
                                if not purpose_has_base_kw:
                                    rec['말소후보'] = False

                            rec['절대인수'] = any(kw in content for kw in always_keep_keywords)
                            rec['AI해석필요'] = any(kw in content for kw in ai_check_keywords)
                            rec['소유권이전'] = '이전' in content and not rec['말소후보'] and not rec['절대인수']

                            purpose_text = rec.get('등기목적', '')
                            malso_purpose_kws = ['말소', '抹消', '취소', '해지', '해제']
                            has_malso_in_purpose = any(mk in purpose_text for mk in malso_purpose_kws)
                            # ⚠️ 등기목적(purpose)에만 말소 키워드가 있을 때만 '이미 말소됨' 처리
                            # content 전체에서 '말소' 검색 시 하단 안내문구 때문에 가짜 말소로 오인됨
                            malso_combined_purpose = ['근저당권말소', '가압류말소', '압류말소', '경매개시결정말소',
                                              '저당권말소', '담보가등기말소', '전세권말소', '근저당말소']
                            has_malso_combined = any(mc in purpose_text for mc in malso_combined_purpose)
                            rec['이미말소됨'] = has_malso_in_purpose or has_malso_combined

                            rec['접수번호_오타'] = ""
                            receipt_match = re.search(r'제\s*(\d+)\s*호', rec['전체내용'])
                            if receipt_match:
                                receipt_num_check = receipt_match.group(1)
                                if len(receipt_num_check) <= 1 or len(receipt_num_check) >= 8:
                                    rec['접수번호_오타'] = f"⚠️ 접수번호 '{receipt_num_check}'이(가) 패턴상 오타로 보입니다. 원본 확인 필요."
                            elif rec.get('접수번호'):
                                # fallback 패턴으로 접수번호가 이미 추출되었으므로 오타 경고 불필요
                                pass
                            elif rec.get('접수일자_표시', '') != '확인불가':
                                rec['접수번호_오타'] = "⚠️ 접수번호(제____호)가 인식되지 않았습니다. 원본 확인 필요."

                            _parsed.append(rec)
                        return _parsed

                    # 1차 시도: 분류된 analysis_rows로 파싱
                    parsed_records = parse_records_from_rows(analysis_rows)

                    # 🛡️ Fallback: 파싱 실패 시 원본 OCR 텍스트로 재시도
                    if not parsed_records and analysis_rows != original_clean_rows:
                        st.toast('🔄 분류 결과로 인식 실패 — 원본 텍스트로 재시도합니다...')
                        parsed_records = parse_records_from_rows(original_clean_rows)

                    # 📊 빈 데이터 방지 (등기부 아닌 사진 업로드 시)
                    if not parsed_records:
                        st.warning("⚠️ 등기부등본 내용을 인식하지 못했습니다. 사진을 확인해 주세요.")
                        st.stop()

                    # 📝 이름 오독 통일 (조충희/조충회 등 → 빈도 최고 이름으로)
                    parsed_records = normalize_names(parsed_records)

                    df = pd.DataFrame(parsed_records)
                    candidates = df[df['말소후보'] == True].dropna(subset=['접수일자_기준'])
                    base_date = candidates.sort_values(by='접수일자_기준').iloc[0]['접수일자_기준'] if not candidates.empty else None

                    # 📅 말소기준권리 정보 저장 (시각적 표시용)
                    base_date_info = None
                    if not candidates.empty:
                        first_candidate = candidates.sort_values(by='접수일자_기준').iloc[0]
                        base_date_info = {
                            'date': base_date,
                            'purpose': first_candidate.get('등기목적', ''),
                            'gu': first_candidate.get('구분', ''),
                            'rank': first_candidate.get('순위번호', ''),
                        }
                    st.session_state.base_date_info = base_date_info

                    def determine_status(row):
                        """auction_rules.txt 기반 하드코딩 판별 — 95%+ 즉시 판별"""
                        content = str(row.get('전체내용', ''))
                        purpose = str(row.get('등기목적', ''))
                        has_date = pd.notnull(row.get('접수일자_기준')) and base_date is not None
                        is_senior = has_date and row['접수일자_기준'] < base_date
                        is_junior = has_date and row['접수일자_기준'] >= base_date

                        # === 0순위: 이미 말소(가로줄) ===
                        if row.get('이미말소됨', False):
                            return "🔘 이미 말소됨", ""

                        # === 1순위: 무조건 인수 (날짜 불문) ===
                        if '유치권' in content:
                            return "🚨 절대 인수", "유치권은 매각과 무관하게 매수인이 무조건 인수해야 합니다."
                        if '법정지상권' in content or '관습상' in content or '관습법상' in content:
                            return "🚨 절대 인수", "법정지상권/관습상 법정지상권은 날짜와 무관하게 무조건 인수됩니다."
                        if '분묘기지권' in content:
                            return "🚨 절대 인수", "분묘기지권은 무조건 인수되며 토지 사용에 심각한 제약이 있습니다."
                        if '예고등기' in content:
                            return "🚨 절대 인수", "예고등기는 소유권 관련 소송이 진행 중임을 나타내며, 무조건 인수됩니다."
                        if '도시철도법' in content or ('구분지상권' in content and any(k in content for k in ['사용재결', '토지수용', '공익', '도시철도'])):
                            return "🚨 절대 인수", "공익사업 목적 구분지상권은 순위 불문 무조건 인수됩니다."
                        if '채무자회생법' in content and '폐지' not in content:
                            return "🚨 절대 인수", "채무자회생법에 의한 등기는 말소촉탁 대상이 아닙니다."
                        if '건물철거' in content or '토지인도' in content:
                            return "🚨 절대 인수", "건물철거/토지인도 목적 가처분은 후순위라도 절대 인수됩니다."
                        if '특별매각조건' in content or '인수조건' in content:
                            return "🚨 절대 인수", "특별매각조건에 따라 매수인 인수가 확정된 권리입니다."
                        if '요역지' in content and '지역권' in content:
                            return "🚨 절대 인수", "요역지 지역권은 소유권에 부종하므로 무조건 인수됩니다."

                        # === 2순위: 무조건 말소 (담보물권/압류성 권리) ===
                        if row.get('말소후보', False):
                            return "❌ 말소", "담보물권/압류성 권리는 매각 시 무조건 소멸합니다."

                        # === 3순위: 소유권이전 등 기본등기 ===
                        if row.get('소유권이전', False):
                            return "➖ 기본등기", ""

                        # === 4순위: 조건부 판단 (auction_rules 하드코딩) ===

                        # 4-1: 전세권 — 배당요구 여부로 분기
                        if '전세권' in purpose or '전세권' in content:
                            if is_junior:
                                return "❌ 말소", "후순위 전세권은 매각으로 소멸합니다."
                            if is_senior:
                                if '배당요구' in content:
                                    return "❌ 말소", "배당요구한 선순위 전세권은 소멸합니다 (일부배당도 전액말소)."
                                return "✅ 인수", "배당요구 없는 선순위 전세권은 전세금 전액을 매수인이 인수합니다."
                            return "🤖 AI 정밀해석", ""

                        # 4-2: 가처분 — 일반 선후 판단
                        if '가처분' in purpose or '처분금지' in purpose or '가처분' in content:
                            if is_senior:
                                return "✅ 인수", "선순위 가처분은 매수인이 인수하며, 소유권 상실 위험이 있습니다."
                            if is_junior:
                                return "❌ 말소", "후순위 가처분은 매각으로 소멸합니다."
                            return "🤖 AI 정밀해석", ""

                        # 4-3: 가등기 — 담보가등기 vs 보전가등기
                        if '가등기' in purpose or '가등기' in content:
                            if '담보가등기' in content or '담보' in purpose:
                                return "❌ 말소", "담보가등기는 저당권과 동일하게 취급되어 무조건 소멸합니다."
                            if is_senior:
                                return "✅ 인수", "선순위 보전가등기는 매수인이 인수하며, 본등기 시 소유권 상실 위험이 있습니다."
                            if is_junior:
                                return "❌ 말소", "후순위 가등기는 매각으로 소멸합니다."
                            return "🤖 AI 정밀해석", ""

                        # 4-4: 지상권 — 담보지상권 vs 일반
                        if '지상권' in purpose or ('지상권' in content and '구분' not in content):
                            if '담보' in content or '담보' in purpose:
                                return "❌ 말소", "담보지상권은 피담보채권(근저당권) 소멸 시 함께 소멸합니다."
                            if is_senior:
                                return "✅ 인수", "선순위 지상권은 매수인이 인수합니다."
                            if is_junior:
                                return "❌ 말소", "후순위 지상권은 매각으로 소멸합니다."
                            return "🤖 AI 정밀해석", ""

                        # 4-5: 지역권
                        if '지역권' in purpose or '지역권' in content:
                            if is_senior:
                                return "✅ 인수", "선순위 지역권은 매수인이 인수합니다."
                            if is_junior:
                                return "❌ 말소", "후순위 지역권은 매각으로 소멸합니다."
                            return "🤖 AI 정밀해석", ""

                        # 4-6: 임차권 — 대항력/배당 판단은 OCR만으로 불가 → AI
                        if '임차권' in purpose or '임차권' in content or '임차인' in content:
                            return "🤖 AI 정밀해석", ""

                        # === 4-7: 동일 날짜 배틀 ===
                        # 전입신고일과 근저당 설정일이 동일하면
                        # 근저당이 우선 (등기는 당일 효력, 전입신고는 다음날 0시 효력)
                        if has_date and row['접수일자_기준'] == base_date:
                            # 임차 관련 권리는 동일일이면 근저당이 우선 (말소)
                            if '임차' in content or '전세' in content or '전입' in content:
                                return "❌ 말소", (
                                    "전입신고일과 말소기준권리 설정일이 동일합니다. "
                                    "전입신고는 다음날 0시에 효력이 발생하나 "
                                    "근저당 등기는 당일 접수 시 즉시 효력이 발생하므로 "
                                    "근저당이 우선하여 임차인은 대항력이 없습니다(말소)."
                                )

                        # === 5순위: 일반 날짜 비교 ===
                        if is_junior:
                            return "❌ 말소", ""
                        if is_senior:
                            return "✅ 인수", ""

                        return "기타", ""

                    # determine_status 적용 (결과 + 하드코딩 이유)
                    df['AI_상세이유'] = ""
                    for idx in df.index:
                        result, reason = determine_status(df.loc[idx])
                        df.at[idx, '결과'] = result
                        if reason:
                            df.at[idx, 'AI_상세이유'] = reason

                # 🤖 Gemini AI 정밀 해석 (묶음 처리 — 1회 API 호출)
                ai_targets = df[df['결과'].str.contains('AI 정밀해석')].index.tolist()
                if ai_targets:
                    with st.spinner(f'🤖 AI 정밀 해석 중... ({len(ai_targets)}건 일괄 처리)'):
                        ai_rows_data = []
                        for index in ai_targets:
                            row = df.loc[index]
                            ai_rows_data.append({
                                'index': int(index),
                                'content': row['전체내용'],
                                'section_gu': row.get('구분', None),
                            })

                        try:
                            batch_results = ask_gemini_for_rights_batch(
                                ai_rows_data, base_date, model, spec_summary
                            )

                            for index in ai_targets:
                                result_data = batch_results.get(int(index), {})
                                result_str = result_data.get('결과', '추가확인')
                                reason = result_data.get('이유', '')

                                if '인수' in result_str:
                                    df.at[index, '결과'] = "✅ 인수 (AI판단)"
                                elif '말소' in result_str:
                                    df.at[index, '결과'] = "❌ 말소 (AI판단)"
                                else:
                                    df.at[index, '결과'] = "⚠️ 서류확인 요망"
                                df.at[index, 'AI_상세이유'] = reason

                        except Exception as e:
                            for index in ai_targets:
                                df.at[index, '결과'] = "⚠️ 판단 지연(수동 확인)"
                                df.at[index, 'AI_상세이유'] = f"AI 일괄 처리 실패: {str(e)[:80]}"

                # =====================================================
                # 🔬 고급 교차 검증 (등기부 전체 + 명세서 교차 분석)
                # =====================================================
                cross_warnings = []

                # 1. 대위변제 위험 감지
                cross_warnings.extend(detect_daewi_risk(df, parsed_records))

                # 2. 조세채권(당해세) 충돌 감지
                cross_warnings.extend(detect_tax_seizure_conflict(df, parsed_records))

                # 3. 매각물건명세서 Override (비고란 키워드로 기존 판단 덮어쓰기)
                cross_warnings.extend(apply_spec_overrides(df, spec_summary))

                # 4. 지분경매 및 신탁등기 감지
                cross_warnings.extend(detect_share_auction_and_trust(df, parsed_records, all_clean_rows))

                # 5. 전 소유자 가압류/가처분 교차 검증
                cross_warnings.extend(detect_prev_owner_claims(df, parsed_records, spec_summary))

                # 6. 최우선 임금채권 경고
                cross_warnings.extend(detect_wage_claim_risk(df, parsed_records))

                # 7. 지분경매 근저당 범위 판별
                cross_warnings.extend(detect_share_mortgage_scope(df, parsed_records, all_clean_rows))

                # 8. 명도 난이도 평가
                cross_warnings.extend(evaluate_eviction_difficulty(df, base_date, parsed_records))

                st.session_state.cross_warnings = cross_warnings

                # 🔍 매각 후 소멸(말소) 대상 권리 분석 수행
                malso_omission_report = None
                if base_date and len(parsed_records) > 0:
                    with st.spinner('🔍 Gemini가 매각 후 소멸 대상 권리를 분석하고 있습니다...'):
                        all_records_text = "\n".join([r['전체내용'] for r in parsed_records])
                        # 확정된 말소 목록을 텍스트로 변환하여 참조 데이터로 전달
                        _confirmed_malso_text = ""
                        try:
                            for _mi, _mr in malso_df.iterrows():
                                _confirmed_malso_text += f"  {_mi}. {_mr['구분']} 순위번호 {_mr['순위번호']}번 - {_mr['등기목적']} ({_mr['접수일자']})\n"
                        except Exception:
                            pass
                        malso_omission_report = ask_gemini_for_malso_omission(
                            all_records_text, base_date, model, spec_summary,
                            confirmed_malso_summary=_confirmed_malso_text if _confirmed_malso_text else None
                        )
                st.session_state.malso_omission_report = malso_omission_report

                # '이미 말소됨' 항목은 말소 목록에서 제외 (이미 처리된 건이므로)
                malso_df = df[
                    df['결과'].str.contains('말소') & ~df['결과'].str.contains('이미 말소됨')
                ][['구분', '순위번호', '등기목적', '접수일자_표시']]
                malso_df.columns = ['구분', '순위번호', '등기목적', '접수일자']
                malso_df.index = range(1, len(malso_df) + 1)

                st.session_state.final_df = df
                st.session_state.malso_df = malso_df

                # 🧾 Feature A: Gemini 종합 안전도 리포트 생성
                safety_report = None
                if base_date:
                    with st.spinner('🧾 Gemini가 종합 안전도를 평가하고 있습니다...'):
                        # 확정된 말소 목록을 텍스트로 변환하여 참조 데이터로 전달
                        _confirmed_malso_text_sr = ""
                        try:
                            for _mi, _mr in malso_df.iterrows():
                                _confirmed_malso_text_sr += f"  {_mi}. {_mr['구분']} 순위번호 {_mr['순위번호']}번 - {_mr['등기목적']} ({_mr['접수일자']})\n"
                        except Exception:
                            pass
                        safety_report = ask_gemini_for_safety_report(
                            df, base_date, model, spec_summary, parsed_records,
                            confirmed_malso_summary=_confirmed_malso_text_sr if _confirmed_malso_text_sr else None
                        )
                st.session_state.safety_report = safety_report



                st.session_state.step = 2
                st.rerun()

            except Exception as e:
                st.error(f"분석 중 오류가 발생했습니다: {e}")

    st.markdown("<br>", unsafe_allow_html=True)
    
    st.info("💡 **최고의 인식률을 위한 꿀팁!**\n\n무료 스캐너 앱 **'vFlat'**으로 문서를 찍어서 올리시면, 사진 용량이 1/10로 줄어들어 분석 속도와 인식률이 비약적으로 상승합니다.\n\n🍎 [아이폰 vFlat 설치](https://apps.apple.com/kr/app/vflat-scan-pdf-scanner/id1540238220) &nbsp;&nbsp;|&nbsp;&nbsp; 🤖 [갤럭시 vFlat 설치](https://play.google.com/store/apps/details?id=com.voyagerx.scanner)")
    st.markdown("<br>", unsafe_allow_html=True)

    # 🌟 수정된 주의사항 (요청하신 대로 군더더기 없이 깔끔하게!)
    with st.expander("🚨 주의사항 및 개인정보 보호 안내 (클릭해서 확인)"):
        st.markdown("""
        * **[면책조항]** AI 판독 결과는 100% 완벽하지 않을 수 있으며, 오류가 발생할 수 있습니다. 본 결과는 참고용으로만 활용하시기 바랍니다.
        * **[개인정보 보호]** 업로드하신 사진은 서버에 일절 저장되지 않습니다. 권리분석 완료 후 즉시 메모리에서 영구 삭제됩니다.
        """)

# =====================================================================
# 📑 [2단계 화면] 결과 및 다운로드
# =====================================================================
elif st.session_state.step == 2:
    st.title("✨ 분석이 완료되었습니다!")

    # =========================================================
    # ① 📊 분석 요약 대시보드 (기존 동일)
    # =========================================================
    if st.session_state.final_df is not None:
        result_df = st.session_state.final_df
        total = len(result_df)
        insu_count = len(result_df[result_df['결과'].str.contains('인수', na=False) & ~result_df['결과'].str.contains('이미|절대', na=False)])
        malso_count = len(result_df[result_df['결과'].str.contains('말소', na=False) & ~result_df['결과'].str.contains('이미', na=False)])
        already_count = len(result_df[result_df['결과'].str.contains('이미 말소됨', na=False)])
        ai_count = len(result_df[result_df['결과'].str.contains('서류확인|AI판단', na=False, regex=True)])
        danger_count = len(result_df[result_df['결과'].str.contains('절대 인수', na=False)])

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("📋 총 등기", f"{total}건")
        col2.metric("✅ 인수", f"{insu_count}건")
        col3.metric("❌ 말소", f"{malso_count}건")
        col4.metric("🔘 이미말소", f"{already_count}건")

        if danger_count > 0 or ai_count > 0:
            col5, col6 = st.columns(2)
            if danger_count > 0:
                col5.metric("🚨 절대 인수", f"{danger_count}건")
            if ai_count > 0:
                col6.metric("🤖 AI 판단", f"{ai_count}건")

        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ② 📅 말소기준권리 일자 시각적 강조
    # =========================================================
    if st.session_state.base_date_info:
        bdi = st.session_state.base_date_info
        st.info(f"📅 **말소기준권리**: {bdi['date']}  |  {bdi['gu']} 순위번호 {bdi['rank']}번  |  {bdi['purpose']}\n\n이 날짜 이후(같은 날 포함)에 접수된 말소 대상 등기는 낙찰로 소멸됩니다.")
        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ③ 🧾 종합 안전도 리포트 (기존 동일)
    # =========================================================
    if st.session_state.safety_report:
        with st.expander("🧾 종합 안전도 리포트 (Gemini 평가)", expanded=True):
            try:
                report = st.session_state.safety_report
                if '🟢' in report:
                    st.success(report)
                elif '🔴' in report:
                    st.error(report)
                else:
                    st.warning(report)
            except Exception as report_err:
                st.warning(f"⚠️ 종합 안전도 리포트 표시 중 오류: {report_err}")
        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ④ 🚨 매각물건명세서 위험 경고
    # =========================================================
    if st.session_state.danger_warnings:
        st.subheader("🚨 매각물건명세서 위험 경고")
        for warning in st.session_state.danger_warnings:
            st.error(warning)
        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ⑤ 📋 매각물건명세서 교차 검증 결과
    # =========================================================
    if st.session_state.spec_summary:
        with st.expander("📋 매각물건명세서 분석 결과 (교차 검증 완료)", expanded=True):
            st.info("📋 매각물건명세서가 감지되어 등기부등본과 교차 검증을 수행했습니다. AI 판단의 정확도가 향상되었습니다.")
            st.markdown(st.session_state.spec_summary)
        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ⑥ 🤖 AI 상세 판독 내역 및 수정 (st.data_editor)
    #    → 다운로드보다 위에 배치하여 수정 후 반영되도록
    # =========================================================
    with st.expander("🤖 AI 상세 판독 내역 및 수정 (클릭하여 결과 직접 교정)", expanded=True):
        display_cols = ['구분', '순위번호', '등기목적', '결과', 'AI_상세이유']
        edited_df = st.data_editor(
            st.session_state.final_df[display_cols],
            use_container_width=True,
            disabled=['구분', '순위번호', '등기목적'],
            key='rights_editor'
        )
        # 수정 내역을 final_df에 즉시 반영
        st.session_state.final_df['결과'] = edited_df['결과'].values
        st.session_state.final_df['AI_상세이유'] = edited_df['AI_상세이유'].values

        # � malso_df도 수정된 final_df 기반으로 재필터링 (데이터 동기화)
        _updated_malso = st.session_state.final_df[
            st.session_state.final_df['결과'].str.contains('말소', na=False) &
            ~st.session_state.final_df['결과'].str.contains('이미 말소됨', na=False)
        ][['구분', '순위번호', '등기목적', '접수일자_표시']]
        _updated_malso.columns = ['구분', '순위번호', '등기목적', '접수일자']
        _updated_malso.index = range(1, len(_updated_malso) + 1)
        st.session_state.malso_df = _updated_malso

        # 📝 접수번호 OCR 오타 경고 표시
        if '접수번호_오타' in st.session_state.final_df.columns:
            typo_rows = st.session_state.final_df[st.session_state.final_df['접수번호_오타'] != ''].copy()
            if not typo_rows.empty:
                st.markdown("---")
                st.markdown("**📝 접수번호 OCR 오타 감지 결과**")
                for _, row in typo_rows.iterrows():
                    st.warning(f"순위번호 {row['순위번호']}번: {row['접수번호_오타']}")

    st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ⑦ 📊 권리 타임라인 시각화 (에디터 수정 후 색상 연동)
    # =========================================================
    if st.session_state.final_df is not None:
        timeline_df = st.session_state.final_df.dropna(subset=['접수일자_기준']).copy()
        if not timeline_df.empty:
            with st.expander("📊 권리 타임라인 시각화", expanded=True):
              try:
                result_color_map = {
                    '인수': '#2ecc71',
                    '말소': '#e74c3c',
                    '절대 인수': '#e67e22',
                    '이미 말소됨': '#95a5a6',
                    'AI판단': '#3498db',
                    '서류확인': '#f39c12',
                    '기본등기': '#bdc3c7',
                    '기타': '#7f8c8d',
                }

                purpose_color_map = {
                    '근저당': '#E74C3C',
                    '저당': '#E74C3C',
                    '압류': '#C0392B',
                    '가압류': '#E67E22',
                    '경매개시결정': '#D35400',
                    '전세권': '#2980B9',
                    '임차권': '#3498DB',
                    '지상권': '#27AE60',
                    '가처분': '#8E44AD',
                    '가등기': '#9B59B6',
                    '소유권이전': '#1ABC9C',
                    '유치권': '#F39C12',
                }

                # 마커 크기를 결과 중요도에 따라 차별화
                def get_marker_size(result_str):
                    r = str(result_str)
                    if '절대 인수' in r:
                        return 22
                    if '인수' in r:
                        return 18
                    if '말소' in r and '이미' not in r:
                        return 16
                    if '이미' in r:
                        return 10
                    if '기본' in r:
                        return 10
                    return 14

                def get_purpose_color(purpose, result):
                    """등기목적 기반 색상 → 결과 기반 색상 fallback"""
                    purpose_clean = str(purpose).replace(' ', '')
                    for key, color in purpose_color_map.items():
                        if key in purpose_clean:
                            return color
                    for key, color in result_color_map.items():
                        if key in str(result):
                            return color
                    return '#7f8c8d'

                timeline_df['색상'] = timeline_df.apply(
                    lambda r: get_purpose_color(r.get('등기목적', ''), r.get('결과', '')), axis=1
                )
                timeline_df['Y축_base'] = timeline_df['구분'].apply(lambda x: 1 if '갑' in x else 2)
                timeline_df['날짜_str'] = timeline_df['접수일자_기준'].apply(lambda d: d.isoformat() if d else None)

                # 동일 날짜 + 동일 구분 마커가 겹치지 않도록 Y축 jitter 적용
                from collections import Counter
                date_gu_counter = Counter()
                y_jitter = []
                for _, row in timeline_df.iterrows():
                    key = (row['날짜_str'], row['Y축_base'])
                    count = date_gu_counter[key]
                    date_gu_counter[key] += 1
                    jitter_offsets = [0, 0.15, -0.15, 0.25, -0.25, 0.35, -0.35]
                    offset = jitter_offsets[count] if count < len(jitter_offsets) else 0.1 * count
                    y_jitter.append(row['Y축_base'] + offset)
                timeline_df['Y축'] = y_jitter

                fig = go.Figure()

                # 말소기준권리 이전/이후 배경색 구분 (반투명 영역)
                if st.session_state.base_date_info:
                    bd = st.session_state.base_date_info['date']
                    bd_str = bd.isoformat() if hasattr(bd, 'isoformat') else str(bd)
                    # 말소기준 이전: 연한 초록 배경 (안전)
                    fig.add_vrect(
                        x0=timeline_df['날짜_str'].min(), x1=bd_str,
                        fillcolor='rgba(46,204,113,0.06)', line_width=0,
                        annotation_text='말소기준 이전 (안전)', annotation_position='top left',
                        annotation_font_size=10, annotation_font_color='#27ae60',
                    )
                    # 말소기준 이후: 연한 빨강 배경 (말소 대상)
                    fig.add_vrect(
                        x0=bd_str, x1=timeline_df['날짜_str'].max(),
                        fillcolor='rgba(231,76,60,0.06)', line_width=0,
                        annotation_text='말소기준 이후 (말소)', annotation_position='top right',
                        annotation_font_size=10, annotation_font_color='#e74c3c',
                    )

                for _, row in timeline_df.iterrows():
                    hover_lines = [
                        f"<b>{row['구분']} #{row['순위번호']}</b>",
                        f"📋 {row['등기목적']}",
                        f"📅 {row.get('접수일자_표시', '')}",
                    ]
                    if row.get('접수번호'):
                        hover_lines.append(f"🔢 접수번호: 제{row['접수번호']}호")
                    content = str(row.get('전체내용', ''))
                    amt_match = re.search(r'(?:채권최고액|금)\s*([\d,]+)\s*원', content.replace(' ', ''))
                    if amt_match:
                        try:
                            amt_val = int(amt_match.group(1).replace(',', ''))
                            hover_lines.append(f"💰 채권최고액: {amt_val:,}원")
                        except ValueError:
                            pass
                    hover_lines.append(f"🏷️ {row['결과']}")
                    if row.get('AI_상세이유'):
                        hover_lines.append(f"💬 {str(row['AI_상세이유'])[:60]}")
                    hover_text = "<br>".join(hover_lines)

                    m_size = get_marker_size(row['결과'])

                    fig.add_trace(go.Scatter(
                        x=[row['날짜_str']],
                        y=[row['Y축']],
                        mode='markers+text',
                        marker=dict(size=m_size, color=row['색상'],
                                    line=dict(width=2, color='white'),
                                    opacity=0.9),
                        text=[str(row['순위번호'])],
                        textposition='top center',
                        textfont=dict(size=10, color='#333'),
                        hovertext=hover_text,
                        hoverinfo='text',
                        showlegend=False,
                    ))

                if st.session_state.base_date_info:
                    bd = st.session_state.base_date_info['date']
                    bd_str = bd.isoformat() if hasattr(bd, 'isoformat') else str(bd)
                    fig.add_shape(
                        type='line',
                        x0=bd_str, x1=bd_str, y0=0.2, y1=2.8,
                        line=dict(color='red', width=3, dash='dash'),
                    )
                    fig.add_annotation(
                        x=bd_str, y=2.9,
                        text='📌 말소기준권리',
                        showarrow=False,
                        font=dict(color='red', size=13, family='Pretendard'),
                        yshift=12,
                    )

                legend_items = [
                    ('근저당/저당', '#E74C3C'), ('압류', '#C0392B'),
                    ('가압류', '#E67E22'), ('경매개시결정', '#D35400'),
                    ('전세권', '#2980B9'), ('임차권', '#3498DB'),
                    ('가처분', '#8E44AD'), ('소유권이전', '#1ABC9C'),
                    ('이미말소', '#95a5a6'),
                ]
                for label, color in legend_items:
                    fig.add_trace(go.Scatter(
                        x=[None], y=[None], mode='markers',
                        marker=dict(size=10, color=color),
                        name=label,
                    ))

                fig.update_layout(
                    title=dict(text='📊 등기 접수일자 타임라인 (권리 종류별 색상)', font=dict(size=16, color='#333')),
                    xaxis_title='접수일자',
                    xaxis=dict(
                        gridcolor='rgba(200,200,200,0.3)',
                        linecolor='#999',
                        tickfont=dict(color='#333', size=11),
                        title_font=dict(color='#333'),
                    ),
                    yaxis=dict(
                        tickvals=[1, 2],
                        ticktext=['갑구 (소유권)', '을구 (기타권리)'],
                        range=[0.2, 3.2],
                        gridcolor='rgba(200,200,200,0.3)',
                        linecolor='#999',
                        tickfont=dict(color='#333', size=12),
                    ),
                    height=600,
                    margin=dict(l=20, r=20, t=60, b=40),
                    legend=dict(orientation='h', yanchor='bottom', y=-0.2, xanchor='center', x=0.5,
                                font=dict(size=11, color='#333')),
                    paper_bgcolor='rgba(253,251,247,1)',
                    plot_bgcolor='rgba(253,251,247,1)',
                    font=dict(color='#333'),
                    hoverlabel=dict(bgcolor='white', font_size=14, font_family='Pretendard',
                                    bordercolor='#ddd'),
                )
                st.plotly_chart(fig, use_container_width=True)
              except Exception as chart_err:
                st.warning(f"⚠️ 타임라인 차트 표시 중 오류: {chart_err}")
            st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # ⑧ 📑 말소할 등기 목록 + � 다운로드 (수정된 malso_df 사용)
    # =========================================================
    st.subheader("📑 법원 제출용: 말소할 등기 목록")
    st.table(st.session_state.malso_df)

    # 📥 DOCX 문서 생성 (최종 업데이트된 malso_df 기반)
    doc = Document()
    doc.add_heading('말 소 할  등 기  목 록', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for idx, row in st.session_state.malso_df.iterrows():
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {row['구분']} 순위번호 제{row['순위번호']}번\n").bold = True
        p.add_run(f"   {row['등기목적']}\n")
        p.add_run(f"   {row['접수일자']} 접수")

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # 📥 PDF 문서 생성
    @st.cache_resource
    def _get_pdf_font_path():
        """PDF용 한글 폰트를 app.py와 같은 폴더의 로컬 파일에서 직접 참조합니다."""
        import os
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'NanumGothic-Regular.ttf')

    def generate_pdf(malso_df):
        try:
            from fpdf import FPDF
            font_path = _get_pdf_font_path()
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font('NanumGothic', '', font_path, uni=True)
            pdf.set_font('NanumGothic', size=16)
            pdf.cell(0, 12, '말 소 할  등 기  목 록', ln=True, align='C')
            pdf.ln(8)
            pdf.set_font('NanumGothic', size=11)
            for idx, row in malso_df.iterrows():
                pdf.cell(0, 8, f"{idx}. {row['구분']} 순위번호 제{row['순위번호']}번", ln=True)
                pdf.cell(0, 7, f"   {row['등기목적']}", ln=True)
                pdf.cell(0, 7, f"   {row['접수일자']} 접수", ln=True)
                pdf.ln(4)
            return bytes(pdf.output())
        except Exception:
            return None

    # 📥 다운로드 버튼 (DOCX + PDF)
    dl_col1, dl_col2 = st.columns(2)
    with dl_col1:
        st.download_button(
            label="📥 워드(.docx) 다운로드",
            data=doc_io,
            file_name="말소할_등기_목록.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    with dl_col2:
        pdf_data = generate_pdf(st.session_state.malso_df)
        if pdf_data:
            st.download_button(
                label="📥 PDF 다운로드",
                data=pdf_data,
                file_name="말소할_등기_목록.pdf",
                mime="application/pdf",
                type="primary",
                use_container_width=True
            )
        else:
            st.caption("PDF 생성 불가 (fpdf2 미설치)")

    st.markdown("<br><hr><br>", unsafe_allow_html=True)

    # =========================================================
    # ⑨ 🔬 고급 교차 검증 경고 (맨 마지막 배치)
    # =========================================================
    cross_warnings = st.session_state.get('cross_warnings', [])
    if cross_warnings:
        with st.expander(f"🔬 고급 교차 검증 경고 ({len(cross_warnings)}건)", expanded=True):
            for w in cross_warnings:
                if '🚨' in w or '⛔' in w:
                    st.error(w)
                elif '❗' in w:
                    st.error(w)
                elif '💡' in w:
                    st.info(w)
                else:
                    st.warning(w)

    # =========================================================
    # ⑩ 🔍 매각 후 소멸(말소) 대상 권리 분석 보고서
    # =========================================================
    if st.session_state.malso_omission_report:
        with st.expander("🔍 매각 후 소멸(말소) 대상 권리 분석 결과", expanded=True):
            report = st.session_state.malso_omission_report
            if "정상적으로 확인" in report or "소멸 대상 권리가 모두" in report:
                st.success(report)
            elif "매각 시 소멸 확정" in report or "인수 주의 권리" in report:
                st.info("📋 아래는 매각 시 소멸되는 권리와 인수 주의 권리를 AI가 분석한 결과입니다.")
                st.markdown(report)
            else:
                st.warning("⚠️ 아래 분석 결과를 확인하고, 원본 등기부등본과 대조해 주세요.")
                st.markdown(report)
        st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # 💬 분석 결과 Q&A (최대 2회 질문)
    # =========================================================
    st.markdown("---")
    st.subheader("💬 분석 결과에 대해 질문하기")
    st.caption("분석 결과에 대해 궁금한 점을 AI에게 질문할 수 있습니다. (최대 2회)")

    # 기존 Q&A 대화 이력 표시
    if st.session_state.qa_history:
        for qa in st.session_state.qa_history:
            with st.chat_message("user"):
                st.markdown(qa['question'])
            with st.chat_message("assistant"):
                st.markdown(qa['answer'])

    if st.session_state.qa_count < 2:
        qa_remaining = 2 - st.session_state.qa_count
        qa_input = st.text_input(
            "질문을 입력하세요",
            placeholder="예: 인수되는 전세권의 보증금은 얼마인가요? / 이 물건에서 가장 주의해야 할 점은?",
            key="qa_input",
            label_visibility="collapsed"
        )
        if st.button(f"💬 질문하기 (남은 횟수: {qa_remaining}회)", use_container_width=True):
            if qa_input and qa_input.strip():
                with st.spinner("🤖 AI가 답변을 준비하고 있습니다..."):
                    try:
                        genai.configure(api_key=GEMINI_API_KEY)
                        qa_model = genai.GenerativeModel('gemini-3-flash-preview')

                        # 분석 결과 요약 구성
                        qa_context_parts = []
                        if st.session_state.final_df is not None:
                            result_df = st.session_state.final_df
                            for _, r in result_df.iterrows():
                                qa_context_parts.append(
                                    f"{r.get('구분','')} 순위 {r.get('순위번호','')}: {r.get('등기목적','')} → {r.get('결과','')} | {r.get('AI_상세이유','')}"
                                )
                        qa_context = "\n".join(qa_context_parts)

                        safety_ctx = st.session_state.safety_report or ""
                        spec_ctx = st.session_state.spec_summary or ""
                        malso_report_ctx = st.session_state.malso_omission_report or ""

                        qa_prompt = f"""
너는 대한민국 법원 경매 권리분석 최고 전문가야.
아래는 이미 완료된 권리분석 결과야. 사용자의 질문에 이 분석 결과를 바탕으로 정확하고 친절하게 답변해 줘.

[분석 결과]
{qa_context}

[종합 안전도 리포트]
{safety_ctx}

[매각물건명세서 분석]
{spec_ctx}

[매각 후 소멸 권리 분석]
{malso_report_ctx}

[사용자 질문]
{qa_input.strip()}

[지시사항]
- 분석 결과에 기반하여 정확하게 답변해 주세요.
- 확실하지 않은 부분은 "원본 등기부등본을 확인해 주세요"라고 안내해 주세요.
- 답변은 한국어로, 간결하면서도 핵심적인 내용을 포함해 주세요.
"""
                        qa_response = qa_model.generate_content(qa_prompt)
                        answer = qa_response.text

                        st.session_state.qa_history.append({
                            'question': qa_input.strip(),
                            'answer': answer
                        })
                        st.session_state.qa_count += 1
                        st.rerun()

                    except Exception as qa_err:
                        st.error(f"⚠️ AI 답변 생성 중 오류가 발생했습니다: {qa_err}")
            else:
                st.warning("질문을 입력해 주세요.")
    else:
        st.success("✅ 질문 2회를 모두 사용하셨습니다.")

    st.markdown("<br>", unsafe_allow_html=True)

    # =========================================================
    # 🔄 처음으로 돌아가기
    # =========================================================
    if st.button("🔄 처음으로 돌아가기", use_container_width=True):
        for key in ['final_df', 'malso_df', 'spec_summary', 'danger_warnings',
                     'malso_omission_report', 'base_date_info', 'safety_report',
                     'cross_warnings', 'user_requests', 'user_request_count',
                     'qa_history', 'qa_count']:
            if key in st.session_state:
                del st.session_state[key]
        st.session_state.step = 1
        st.rerun()